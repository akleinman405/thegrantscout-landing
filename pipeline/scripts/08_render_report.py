#!/usr/bin/env python3
"""
Script 08: Render Report

Final step - renders the report data into a polished Markdown document
and optionally converts it to a branded Word document for delivery.

Usage:
    python scripts/08_render_report.py --client "PSMF"
    python scripts/08_render_report.py --client "PSMF" --format md
    python scripts/08_render_report.py --client "PSMF" --format docx
    python scripts/08_render_report.py --client "PSMF" --format both
"""

import argparse
import json
import sys
from datetime import date, timedelta
from pathlib import Path
from typing import Dict, List

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))
from scripts.utils.paths import (
    TEMPLATES_DIR,
    get_run_dir,
    get_output_dir,
    load_client_registry,
    setup_logging,
)

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def load_json(path: Path) -> dict:
    """Load JSON file."""
    with open(path) as f:
        return json.load(f)


def save_json(data: dict, path: Path):
    """Save data to JSON file."""
    with open(path, "w") as f:
        json.dump(data, f, indent=2)


# ============================================================================
# Formatting helper functions
# ============================================================================

def format_currency(amount: float) -> str:
    """Format as currency: $1.5M, $50K, $500"""
    if amount is None or amount == 0:
        return "$0"
    if amount >= 1000000:
        return f"${amount/1000000:.1f}M"
    elif amount >= 1000:
        return f"${amount/1000:.0f}K"
    else:
        return f"${amount:,.0f}"


def format_currency_full(amount: float) -> str:
    """Format as full currency: $1,234,567"""
    if amount is None:
        return "$0"
    return f"${amount:,.0f}"


def format_percent(value: float) -> str:
    """Format as percentage: 45%"""
    if value is None:
        return "0%"
    if value <= 1:
        return f"{value:.0%}"
    return f"{value:.0f}%"


def format_date(date_str: str) -> str:
    """Format date: January 6, 2025"""
    if not date_str:
        return ""
    try:
        d = date.fromisoformat(date_str)
        return d.strftime("%B %d, %Y")
    except (ValueError, TypeError):
        return date_str


def format_date_short(date_str: str) -> str:
    """Format date: Jan 6"""
    if not date_str:
        return ""
    try:
        d = date.fromisoformat(date_str)
        return d.strftime("%b %d")
    except (ValueError, TypeError):
        return date_str


def format_month(date_str: str) -> str:
    """Format as month and year: January 2025"""
    if not date_str:
        return ""
    try:
        d = date.fromisoformat(date_str)
        return d.strftime("%B %Y")
    except (ValueError, TypeError):
        return date_str


def status_badge(status: str) -> str:
    """Return emoji badge for status."""
    badges = {
        "HIGH": "HIGH PRIORITY",
        "MEDIUM": "MEDIUM PRIORITY",
        "LOW": "LOW PRIORITY"
    }
    return badges.get(status, status)


def effort_indicator(effort: str) -> str:
    """Return effort indicator."""
    indicators = {
        "Low": "Low Effort",
        "Medium": "Medium Effort",
        "High": "High Effort"
    }
    return indicators.get(effort, effort)


def check_funding_alignment(report_data: dict) -> list:
    """
    Check if funding potential aligns with client's target.
    Returns list of warnings.
    """
    warnings = []

    client = report_data.get("client", {})
    exec_summary = report_data.get("executive_summary", {})
    scenarios = exec_summary.get("funding_scenarios", {})

    # Parse client's target
    target_str = client.get("grant_size_target", "")
    target_min = 0
    if "$25,000" in target_str or "25,000" in target_str:
        target_min = 25000
    elif "$10,000" in target_str or "10,000" in target_str:
        target_min = 10000
    elif "$50,000" in target_str or "50,000" in target_str:
        target_min = 50000
    elif "$100,000" in target_str or "100,000" in target_str:
        target_min = 100000

    if target_min:
        ambitious = scenarios.get("ambitious", 0)
        moderate = scenarios.get("moderate", 0)

        if ambitious < target_min:
            warnings.append(
                f"Best-case funding (${ambitious:,}) is below client's minimum target (${target_min:,})"
            )
        elif moderate < target_min * 0.5:
            warnings.append(
                f"Moderate scenario (${moderate:,}) is well below client's target (${target_min:,})"
            )

    # Check opportunity count
    opportunities = report_data.get("opportunities", [])
    if len(opportunities) < 5:
        warnings.append(f"Only {len(opportunities)} opportunities found (expected 5)")

    # Check for any same-state matches
    same_state_count = sum(1 for o in opportunities if o.get("same_state"))
    client_state = client.get("state", "")
    if client_state and same_state_count == 0:
        warnings.append(f"No {client_state}-based foundations in recommendations")

    return warnings


# ============================================================================
# Rendering functions
# ============================================================================

def render_opportunity(opp: dict, template: str) -> str:
    """Render a single opportunity section."""
    snapshot = opp.get("funder_snapshot", {})
    annual = snapshot.get("annual_giving", {})
    typical = snapshot.get("typical_grant", {})
    geographic = snapshot.get("geographic_focus", {})
    repeat = snapshot.get("repeat_funding", {})
    style = snapshot.get("giving_style", {})
    profile = snapshot.get("recipient_profile", {})
    trend = snapshot.get("funding_trend", {})
    comparable = opp.get("comparable_grant") or {}
    contact = snapshot.get("contact", {})
    officers = snapshot.get("officers", [])

    # Format connections section
    connections_md = ""
    for conn in opp.get("connections", []):
        strength_icon = "**Strong:**" if conn.get("strength") == "strong" else "**Moderate:**"
        connections_md += f"- {strength_icon} {conn.get('description', '')}\n"
    if not connections_md:
        connections_md = "- No direct connections identified yet\n"

    # Format next steps section
    next_steps_md = ""
    for step in opp.get("next_steps", []):
        deadline = format_date_short(step.get("deadline", ""))
        next_steps_md += f"| {step.get('action', '')} | {deadline} | {step.get('owner', 'TBD')} |\n"

    # Format comparable grant
    if comparable and comparable.get("recipient"):
        comparable_md = (
            f"**{comparable.get('recipient', '')}** received "
            f"{format_currency_full(comparable.get('amount', 0))} ({comparable.get('year', '')}) "
            f"for: _{comparable.get('purpose', 'general support')}_"
        )
    else:
        comparable_md = "No directly comparable grant found for your specific profile."

    # Format requirements (placeholder - not populated yet)
    requirements_md = ""
    for req in opp.get("application_requirements", []):
        requirements_md += f"- {req}\n"
    if not requirements_md:
        requirements_md = "- Contact foundation for current application requirements\n"

    # Format officers list
    officers_md = ""
    if officers:
        for officer in officers[:5]:  # Limit to 5
            name = officer.get("name", "")
            title = officer.get("title", "")
            if name and title:
                officers_md += f"- **{name}** - {title}\n"
            elif name:
                officers_md += f"- {name}\n"
    if not officers_md:
        officers_md = "- Officer information not available from 990-PF\n"

    # Format how to apply section
    how_to_apply_md = ""
    if contact.get("how_to_apply"):
        how_to_apply_md = contact["how_to_apply"]
    elif contact.get("restrictions"):
        how_to_apply_md = contact["restrictions"]
    else:
        how_to_apply_md = "Contact the foundation directly for application procedures."

    # Add deadline info if available
    if contact.get("deadlines") and contact["deadlines"].upper() not in ["NONE", "N/A", "SEE ABOVE"]:
        how_to_apply_md += f"\n\n**Deadlines:** {contact['deadlines']}"

    # Format application contact
    app_contact_parts = []
    if contact.get("app_contact_name"):
        app_contact_parts.append(contact["app_contact_name"])
    if contact.get("app_contact_phone"):
        app_contact_parts.append(contact["app_contact_phone"])
    if contact.get("app_contact_email"):
        app_contact_parts.append(contact["app_contact_email"])
    app_contact_str = " | ".join(app_contact_parts) if app_contact_parts else "See website"

    # Build context dict
    ctx = {
        "rank": opp.get("rank", 0),
        "foundation_name": opp.get("foundation_name", ""),
        "foundation_ein": opp.get("foundation_ein", ""),
        "foundation_state": opp.get("foundation_state", ""),
        "status_badge": status_badge(opp.get("status", "")),
        "fit_score": opp.get("fit_score", 0),
        "effort": effort_indicator(opp.get("effort", "")),
        "deadline": opp.get("deadline", "Rolling"),
        "amount_range": f"{format_currency(opp.get('amount_min', 0))} - {format_currency(opp.get('amount_max', 0))}",
        "match_score": opp.get("match_score", 0),
        "same_state": "Yes" if opp.get("same_state") else "No",

        # Snapshot data
        "annual_giving": format_currency(annual.get("total", 0)),
        "annual_grant_count": annual.get("count", 0),
        "annual_year": annual.get("year", ""),
        "typical_median": format_currency(typical.get("median", 0)),
        "typical_range": f"{format_currency(typical.get('min', 0))} - {format_currency(typical.get('max', 0))}",
        "geographic_top": f"{geographic.get('top_state', 'N/A')} ({format_percent(geographic.get('top_state_pct', 0))})",
        "geographic_in_state": format_percent(geographic.get("in_state_pct", 0)),
        "repeat_rate": format_percent(repeat.get("rate", 0)),
        "style_general": format_percent(style.get("general_pct", 0)),
        "style_program": format_percent(style.get("program_pct", 0)),
        "recipient_budget_range": f"{format_currency(profile.get('budget_min', 0))} - {format_currency(profile.get('budget_max', 0))}",
        "recipient_sector": profile.get("primary_sector", "N/A"),
        "trend_direction": trend.get("direction", "Stable"),
        "trend_change": f"{trend.get('change_pct', 0):+.0%}" if trend.get('change_pct') else "+0%",

        # Narratives
        "why_this_fits": opp.get("why_this_fits", ""),
        "positioning_strategy": opp.get("positioning_strategy", ""),

        # Formatted sections
        "connections_section": connections_md,
        "comparable_grant": comparable_md,
        "requirements_section": requirements_md,
        "next_steps_rows": next_steps_md,

        # Contact info from 990-PF
        "website_url": contact.get("website_url") or "Not available",
        "phone": contact.get("phone") or "Not available",
        "app_contact": app_contact_str,
        "officers_list": officers_md,
        "how_to_apply": how_to_apply_md,

        # Legacy contact info (for backwards compatibility)
        "contact_name": contact.get("app_contact_name") or opp.get("contact_name", "Not available"),
        "contact_email": contact.get("app_contact_email") or opp.get("contact_email", ""),
        "contact_phone": contact.get("phone") or opp.get("contact_phone", ""),
        "portal_url": contact.get("website_url") or opp.get("portal_url", "")
    }

    try:
        return template.format(**ctx)
    except KeyError as e:
        print(f"  Warning: Missing template variable: {e}")
        # Return a simplified version
        return f"### #{ctx['rank']}: {ctx['foundation_name']}\n\n{ctx['why_this_fits']}\n\n"


def render_executive_summary(exec_sum: dict, opportunities: list) -> dict:
    """Render executive summary section."""

    # Key strengths as bullets
    strengths_md = ""
    for strength in exec_sum.get("key_strengths", []):
        strengths_md += f"- {strength}\n"

    # Funding scenarios
    scenarios = exec_sum.get("funding_scenarios", {})

    # Top 5 quick reference table
    top5_table = "| Rank | Foundation | Status | Fit | Amount Range |\n"
    top5_table += "|------|------------|--------|-----|-------------|\n"
    for opp in opportunities:
        top5_table += (
            f"| {opp['rank']} | {opp['foundation_name'][:35]} | {opp['status']} | "
            f"{opp['fit_score']}/10 | {format_currency(opp['amount_min'])} - "
            f"{format_currency(opp['amount_max'])} |\n"
        )

    # Urgent actions
    urgent_md = ""
    for action in exec_sum.get("urgent_actions", [])[:3]:
        urgent_md += (
            f"- **{action.get('foundation', '')[:30]}:** {action.get('action', '')} "
            f"(by {format_date_short(action.get('deadline', ''))})\n"
        )

    return {
        "key_strengths": strengths_md,
        "one_thing": exec_sum.get("one_thing", ""),
        "top5_table": top5_table,
        "urgent_actions": urgent_md,
        "conservative": format_currency(scenarios.get("conservative", 0)),
        "moderate": format_currency(scenarios.get("moderate", 0)),
        "ambitious": format_currency(scenarios.get("ambitious", 0))
    }


def render_timeline(timeline: list) -> str:
    """Render 8-week timeline as table."""
    md = "| Week | Dates | Focus | Key Tasks |\n"
    md += "|------|-------|-------|----------|\n"

    for week in timeline:
        week_num = week.get("week", 0)
        dates = f"{format_date_short(week.get('date_start', ''))} - {format_date_short(week.get('date_end', ''))}"
        focus = week.get("focus", "")

        # Summarize tasks
        tasks = week.get("tasks", [])
        if tasks:
            task_summary = "; ".join([t.get("action", "")[:40] for t in tasks[:2]])
            if len(tasks) > 2:
                task_summary += f" (+{len(tasks)-2} more)"
        else:
            task_summary = "—"

        md += f"| {week_num} | {dates} | {focus} | {task_summary} |\n"

    return md


def render_contacts_table(opportunities: list) -> str:
    """Render quick reference contacts table."""
    md = "| Foundation | Phone | Website |\n"
    md += "|------------|-------|--------|\n"

    for opp in opportunities:
        name = opp.get("foundation_name", "")[:30]
        snapshot = opp.get("funder_snapshot", {})
        contact = snapshot.get("contact", {})
        phone = contact.get("phone") or "N/A"
        website = contact.get("website_url")
        if website:
            # Shorten URL for display
            display_url = website.replace("https://", "").replace("http://", "")[:35]
            website_md = f"[{display_url}]({website})"
        else:
            website_md = "N/A"
        md += f"| {name} | {phone} | {website_md} |\n"

    return md


def render_portals_table(opportunities: list) -> str:
    """Render quick reference portals and deadlines table."""
    md = "| Foundation | Deadline | How to Apply |\n"
    md += "|------------|----------|-------------|\n"

    for opp in opportunities:
        name = opp.get("foundation_name", "")[:30]
        snapshot = opp.get("funder_snapshot", {})
        contact = snapshot.get("contact", {})

        # Get deadline from contact data or opportunity
        deadline = contact.get("deadlines") or opp.get("deadline", "Rolling")
        if deadline and deadline.upper() in ["NONE", "N/A"]:
            deadline = "Rolling"

        # Get portal/website
        website = contact.get("website_url")
        if website:
            portal = f"[Website]({website})"
        else:
            portal = "Contact foundation"
        md += f"| {name} | {deadline} | {portal} |\n"

    return md


def generate_anchor(name: str) -> str:
    """Generate a markdown anchor from a foundation name."""
    # Convert to lowercase, replace spaces with hyphens, remove special chars
    anchor = name.lower()
    anchor = anchor.replace(" ", "-")
    anchor = "".join(c for c in anchor if c.isalnum() or c == "-")
    return anchor


def generate_table_of_contents(opportunities: list) -> str:
    """Generate dynamic table of contents with opportunity names.

    Structure:
    1. If You Only Do One Thing
    2. Executive Summary
       - Urgent Actions
       - Funding Scenarios
       - Key Strengths
    3. This Month's Top 5
       - Foundation Name 1
       - Foundation Name 2
       ... (all 5 foundations as subbullets)
    4. 8-Week Action Timeline
    5. Quick Reference
       - Contacts
       - Portals & Deadlines
    6. Next Steps
    """
    toc_lines = []
    item_num = 1

    # Fixed sections at the start
    toc_lines.append(f"{item_num}. If You Only Do One Thing")
    item_num += 1

    toc_lines.append(f"{item_num}. Executive Summary")
    toc_lines.append("   - Urgent Actions")
    toc_lines.append("   - Funding Scenarios")
    toc_lines.append("   - Key Strengths")
    item_num += 1

    # This Month's Top 5 with subbullets listing each foundation (no # symbol)
    toc_lines.append(f"{item_num}. This Month's Top 5")
    for opp in opportunities:
        name = opp.get("foundation_name", "Unknown")
        toc_lines.append(f"   - {name}")
    item_num += 1

    # Fixed sections at the end
    toc_lines.append(f"{item_num}. 8-Week Action Timeline")
    item_num += 1

    toc_lines.append(f"{item_num}. Quick Reference")
    toc_lines.append("   - Contacts")
    toc_lines.append("   - Portals & Deadlines")
    item_num += 1

    toc_lines.append(f"{item_num}. Next Steps")

    return "\n".join(toc_lines)


def render_markdown(report_data: dict, report_template: str, opportunity_template: str) -> str:
    """Render complete markdown report."""

    client = report_data.get("client", {})
    meta = report_data.get("report_meta", {})
    opportunities = report_data.get("opportunities", [])
    exec_sum = report_data.get("executive_summary", {})
    timeline = report_data.get("timeline", [])

    # Generate dynamic table of contents
    toc_md = generate_table_of_contents(opportunities)

    # Render executive summary parts
    exec_parts = render_executive_summary(exec_sum, opportunities)

    # Render each opportunity
    opp_sections = ""
    for opp in opportunities:
        opp_sections += render_opportunity(opp, opportunity_template)
        opp_sections += "\n---\n\n"

    # Render timeline
    timeline_md = render_timeline(timeline)

    # Render quick reference tables
    contacts_md = render_contacts_table(opportunities)
    portals_md = render_portals_table(opportunities)

    # Calculate next report date
    report_date_str = meta.get("report_date", date.today().isoformat())
    try:
        report_date_obj = date.fromisoformat(report_date_str)
        next_report_date = format_date((report_date_obj + timedelta(days=7)).isoformat())
    except (ValueError, TypeError):
        next_report_date = "Next week"

    # Build main context
    ctx = {
        "client_name": client.get("organization_name", ""),
        "week_number": meta.get("week_number", 1),
        "report_month": format_month(report_date_str),
        "report_date": format_date(report_date_str),
        "date_range": f"{format_date_short(meta.get('date_range_start', ''))} - {format_date_short(meta.get('date_range_end', ''))}",

        # Table of contents (dynamically generated)
        "table_of_contents": toc_md,

        # Executive summary
        "one_thing": exec_parts["one_thing"],
        "key_strengths": exec_parts["key_strengths"],
        "top5_table": exec_parts["top5_table"],
        "urgent_actions": exec_parts["urgent_actions"],
        "conservative": exec_parts["conservative"],
        "moderate": exec_parts["moderate"],
        "ambitious": exec_parts["ambitious"],

        # Main content
        "opportunity_sections": opp_sections,
        "timeline_table": timeline_md,

        # Quick reference
        "contacts_table": contacts_md,
        "portals_table": portals_md,

        # Footer
        "generated_date": format_date(report_date_str),
        "next_report_date": next_report_date
    }

    try:
        return report_template.format(**ctx)
    except KeyError as e:
        print(f"Warning: Missing template variable in report_template: {e}")
        return f"# Report for {ctx['client_name']}\n\n{opp_sections}"


def convert_to_docx(markdown_content: str, output_path: str, client_name: str):
    """Convert markdown to Word document with branding."""
    if not HAS_DOCX:
        print("Warning: python-docx not installed. Skipping Word document generation.")
        return False

    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "TheGrantScout | Grant Opportunities Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Confidential - Prepared for {client_name}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown and add content
    lines = markdown_content.split('\n')

    for line in lines:
        line_stripped = line.strip()

        if not line_stripped:
            continue

        # Headers
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)

        # Horizontal rule
        elif line_stripped == '---':
            doc.add_paragraph('_' * 50)

        # Bullet points
        elif line_stripped.startswith('- '):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')

        # Numbered lists
        elif line_stripped and line_stripped[0].isdigit() and '. ' in line_stripped[:4]:
            text = line_stripped.split('. ', 1)[1] if '. ' in line_stripped else line_stripped
            doc.add_paragraph(text, style='List Number')

        # Tables (simplified - just add as paragraphs)
        elif line_stripped.startswith('|'):
            # Skip table formatting rows
            if '---' in line_stripped:
                continue
            # Clean up table row
            cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
            doc.add_paragraph(' | '.join(cells))

        # Regular paragraphs
        else:
            # Handle bold and italic
            if '**' in line_stripped or '_' in line_stripped:
                p = doc.add_paragraph()
                # Simple bold handling
                parts = line_stripped.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd parts are bold
                            run.bold = True
            else:
                doc.add_paragraph(line_stripped)

    doc.save(output_path)
    return True


def update_manifest(output_dir: Path, entry: dict):
    """Update or create manifest.json with new entry."""
    manifest_path = output_dir / "manifest.json"

    if manifest_path.exists():
        manifest = load_json(manifest_path)
    else:
        manifest = {"reports": []}

    # Check if entry already exists (by name and run_date)
    existing_idx = None
    for i, r in enumerate(manifest.get("reports", [])):
        if r.get("name") == entry.get("name") and r.get("run_date") == entry.get("run_date"):
            existing_idx = i
            break

    if existing_idx is not None:
        manifest["reports"][existing_idx] = entry
    else:
        manifest["reports"].append(entry)

    save_json(manifest, manifest_path)


def main():
    parser = argparse.ArgumentParser(
        description="Render report data into Markdown and Word documents"
    )
    parser.add_argument("--client", required=True, help="Client name or short_name")
    parser.add_argument(
        "--format", default="both", choices=["md", "docx", "both"],
        help="Output format: md, docx, or both"
    )
    parser.add_argument(
        "--output", default=None,
        help="Custom output directory (default: outputs/YYYY-MM/)"
    )
    parser.add_argument(
        "--date", default=date.today().isoformat(),
        help="Run date (YYYY-MM-DD)"
    )
    args = parser.parse_args()

    # Find client
    registry = load_client_registry()
    client_entry = None
    for c in registry.get("clients", []):
        if (c["name"].lower() == args.client.lower() or
            c.get("short_name", "").lower() == args.client.lower()):
            client_entry = c
            break

    if not client_entry:
        print(f"ERROR: Client '{args.client}' not found in registry")
        sys.exit(1)

    client_name = client_entry["name"].replace(" ", "_")

    # Set up logging
    logger = setup_logging(client_name, "08_render_report")
    logger.info(f"[08_render_report] Starting for {client_entry['name']}")

    # Find run directory
    run_dir = get_run_dir(client_name, args.date)
    if not run_dir.exists():
        print(f"ERROR: Run directory not found: {run_dir}")
        sys.exit(1)

    # Load report data
    report_data_file = run_dir / "07_report_data.json"
    if not report_data_file.exists():
        print(f"ERROR: Report data not found: {report_data_file}")
        sys.exit(1)

    report_data = load_json(report_data_file)
    logger.info(f"Loaded report data ({len(report_data.get('opportunities', []))} opportunities)")

    print(f"\nRendering report...")
    print(f"Client: {report_data.get('client', {}).get('organization_name', 'Unknown')}")
    print(f"Opportunities: {len(report_data.get('opportunities', []))}")

    # Quality gate - check if results are reasonable
    print("\nChecking report quality...")
    quality_warnings = check_funding_alignment(report_data)

    if quality_warnings:
        print("\n" + "="*60)
        print("QUALITY GATE WARNING")
        print("="*60)
        for w in quality_warnings:
            print(f"  - {w}")
        print("\nThe matched foundations may not align with client needs.")
        print("Consider:")
        print("  1. Re-running Script 02 with --top-k 200 (score more foundations)")
        print("  2. Adjusting filters in Script 05")
        print("  3. Manual review before sending to client")
        print("="*60 + "\n")

        for w in quality_warnings:
            logger.warning(f"Quality gate: {w}")
    else:
        print("Quality check: PASSED")

    # Load templates
    report_template_path = TEMPLATES_DIR / "report_template.md"
    opportunity_template_path = TEMPLATES_DIR / "opportunity_template.md"

    if not report_template_path.exists():
        print(f"ERROR: Report template not found: {report_template_path}")
        sys.exit(1)

    if not opportunity_template_path.exists():
        print(f"ERROR: Opportunity template not found: {opportunity_template_path}")
        sys.exit(1)

    report_template = report_template_path.read_text()
    opportunity_template = opportunity_template_path.read_text()
    logger.info("Loaded templates")

    # Render markdown
    markdown_content = render_markdown(report_data, report_template, opportunity_template)
    logger.info(f"Rendered markdown ({len(markdown_content)} characters)")
    print(f"Rendered markdown: {len(markdown_content)} characters")

    # Determine output paths
    client = report_data.get("client", {})
    meta = report_data.get("report_meta", {})

    short_name = client.get("short_name") or client.get("organization_name", "Report")
    safe_name = "".join(c if c.isalnum() else "_" for c in short_name)
    week = meta.get("week_number", 1)
    report_date = meta.get("report_date", date.today().isoformat())

    # Format: PSMF_2025-12 or PSMF_2025-12-30 for the filename
    report_month_short = report_date[:7]  # YYYY-MM
    filename_base = f"{safe_name}_{report_month_short}"

    # Save to runs folder
    md_run_path = run_dir / "08_report.md"
    md_run_path.write_text(markdown_content)
    logger.info(f"Saved: {md_run_path}")
    print(f"Saved: {md_run_path}")

    # Determine output directory
    if args.output:
        output_dir = Path(args.output)
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = get_output_dir()

    output_format = args.format

    # Save markdown to outputs
    if output_format in ["md", "both"]:
        md_output_path = output_dir / f"{filename_base}.md"
        md_output_path.write_text(markdown_content)
        logger.info(f"Saved: {md_output_path}")
        print(f"Saved: {md_output_path}")

    # Generate Word document
    if output_format in ["docx", "both"]:
        if HAS_DOCX:
            # Save to outputs folder
            docx_output_path = output_dir / f"{filename_base}.docx"
            if convert_to_docx(markdown_content, str(docx_output_path), client.get("organization_name", "")):
                logger.info(f"Saved: {docx_output_path}")
                print(f"Saved: {docx_output_path}")

            # Also save to run folder
            docx_run_path = run_dir / "08_report.docx"
            convert_to_docx(markdown_content, str(docx_run_path), client.get("organization_name", ""))
            logger.info(f"Saved: {docx_run_path}")
            print(f"Saved: {docx_run_path}")
        else:
            print("Warning: python-docx not installed. Install with: pip install python-docx")

    # Update manifest
    update_manifest(output_dir, {
        "name": short_name,
        "client_name": client.get("organization_name", ""),
        "status": "generated",
        "run_date": report_date,
        "week_number": week,
        "report_file_md": f"{filename_base}.md",
        "report_file_docx": f"{filename_base}.docx" if output_format in ["docx", "both"] else None,
        "delivered_date": None
    })
    logger.info("Updated manifest.json")

    # Summary
    print(f"\n{'='*60}")
    print("Report generation complete!")
    print(f"{'='*60}")
    print(f"Client: {client.get('organization_name', 'Unknown')}")
    print(f"Week: {week}")
    print(f"Run folder: {run_dir}")
    print(f"Output folder: {output_dir}")
    print(f"\nFiles generated:")
    print(f"  - {md_run_path.name}")
    if output_format in ["docx", "both"] and HAS_DOCX:
        print(f"  - 08_report.docx")
    print(f"  - {filename_base}.md (in outputs)")
    if output_format in ["docx", "both"] and HAS_DOCX:
        print(f"  - {filename_base}.docx (in outputs)")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
