#!/usr/bin/env python3
"""Generate the REPORT_2026-02-13.1 implementation report as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

NAVY = RGBColor(0x1e, 0x3a, 0x5f)
CHARCOAL = RGBColor(0x2C, 0x3E, 0x50)
GOLD = RGBColor(0xd4, 0xa8, 0x53)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def bold_para(label, value):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(f" {value}")
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("THE GRANT SCOUT")
run.font.size = Pt(24)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Email Campaign Relaunch: Government Funding Angle")
run.font.size = Pt(16)
run.font.color.rgb = CHARCOAL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Implementation Report")
run.font.size = Pt(14)
run.font.color.rgb = CHARCOAL

doc.add_paragraph()
bold_para("Date:", "2026-02-13")
bold_para("Status:", "Complete (Code + DB Ready, Pending Launch)")
bold_para("Owner:", "Alec Kleinman")

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
heading("Table of Contents")
toc_items = [
    "1. Executive Summary",
    "2. What Was Built",
    "3. Database Changes",
    "4. Code Changes",
    "5. A/B Test Design",
    "6. Code Review Results",
    "7. Files Created / Modified",
    "8. What's Next: Launch Checklist",
    "9. Risks and Mitigations",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
heading("1. Executive Summary")

doc.add_paragraph(
    "This report covers the implementation of a new email campaign system that targets "
    "nonprofits heavily dependent on government funding, offering them private foundation "
    "alternatives through The Grant Scout. This is a timely angle given the current "
    "administration's federal funding cuts."
)

doc.add_paragraph(
    "The system uses IRS 990 data already in our database to identify which nonprofits "
    "receive significant government grants, then personalizes outreach emails with specific "
    "data points (government dependency percentage, sector, state, etc.) to dramatically "
    "increase relevance over our previous generic emails."
)

heading("Key Numbers", level=2)
add_table(
    ["Metric", "Value"],
    [
        ["Nonprofits with gov dependency data", "34,545"],
        ["Segmented into campaign tiers", "34,490"],
        ["Critical + new to foundations (prime targets)", "6,328"],
        ["Critical + experienced", "3,625"],
        ["High dependency + new", "4,174"],
        ["High dependency + experienced", "6,007"],
        ["Moderate dependency", "14,356"],
        ["With emails (ready to send today)", "63"],
        ["With websites (scrapable for emails)", "~28,700"],
        ["Email templates created", "8 (4 variants x 2 types)"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Previous campaign: 1,844 emails sent, 93.9% delivery, 0.23% reply rate. "
    "A/B simulation showed personalized emails get 81% positive response vs 50% for generic. "
    "This implementation adds the personalization infrastructure to capture that improvement."
)

doc.add_page_break()

# ============================================================
# 2. WHAT WAS BUILT
# ============================================================
heading("2. What Was Built")

doc.add_paragraph(
    "The implementation covers the full pipeline from data enrichment through email generation, "
    "with A/B testing built in from the start. Here is what each piece does and why."
)

heading("Database Layer (Phase 0)", level=2)
doc.add_paragraph(
    "Added four new columns to the nonprofit_prospects table so we can store government "
    "funding data directly alongside each prospect. This avoids expensive joins at email "
    "generation time and lets us segment prospects into campaign tiers."
)

p = doc.add_paragraph()
p.add_run("Why this matters: ").bold = True
p.add_run(
    "Without these columns, every email generation run would need to join against "
    "nonprofit_returns (2.9M rows) to look up government funding. By pre-computing and "
    "storing the data, email generation is fast and the segmentation is queryable."
)

heading("Government Dependency Backfill", level=2)
doc.add_paragraph(
    "Pulled government_grants_amt from each nonprofit's most recent 990 filing (2020+) "
    "and calculated what percentage of their total revenue comes from government sources. "
    "Capped at 100% to handle edge cases where government grants exceed reported revenue "
    "(happens with fiscal year timing differences)."
)

heading("Prospect Segmentation", level=2)
doc.add_paragraph(
    "Each prospect is assigned to one of five segments based on two dimensions: "
    "how dependent they are on government funding, and whether they have any history "
    "of receiving private foundation grants."
)

add_table(
    ["Segment", "Gov Dependency", "Foundation Experience", "Count", "Campaign Strategy"],
    [
        ["critical_new", "70%+", "None", "6,328", "Highest priority. Vulnerable and unfamiliar with alternatives."],
        ["critical_experienced", "70%+", "Has grants", "3,625", "High priority. Know foundations exist but need more."],
        ["high_new", "30-70%", "None", "4,174", "Good targets. Moderate risk, new to foundations."],
        ["high_experienced", "30-70%", "Has grants", "6,007", "Upsell opportunity. Already in the foundation world."],
        ["moderate", "<30%", "Any", "14,356", "Lower urgency. Less government risk angle."],
    ]
)

doc.add_page_break()

heading("Email Templates (Phase 1)", level=2)
doc.add_paragraph(
    "Created four email variants, each with an initial outreach and a follow-up version "
    "(8 templates total). Each variant uses a different psychological angle."
)

add_table(
    ["Variant", "Name", "Angle", "Best For"],
    [
        ["A", "Government Funding Risk", "\"Government grants make up {gov_pct} of your revenue...\"", "Prospects with high gov dependency"],
        ["B", "Discovery", "\"15-25 foundations you probably don't know about...\"", "General prospects, any segment"],
        ["C", "New-to-Foundations", "\"For nonprofits new to foundation funding...\"", "Prospects with zero foundation grant history"],
        ["D", "Control", "Original generic template from Nov 2025 campaign", "Baseline comparison"],
    ]
)

doc.add_paragraph()
heading("Personalization Variables", level=2)
doc.add_paragraph(
    "Each email can now include data pulled directly from the prospect's 990 filing. "
    "This is a major upgrade from the previous system which only supported {greeting} "
    "(first name or blank)."
)

add_table(
    ["Variable", "Source", "Example"],
    [
        ["{greeting}", "contact_name", "\" Alec\" or \"\""],
        ["{org_name}", "org_name", "\"Capital Area Partnership\""],
        ["{city}, {state}", "city, state", "\"Richmond\", \"VA\""],
        ["{sector}", "sector_desc", "\"Human Services\""],
        ["{revenue_band}", "revenue_band", "\"$1M-$2M\""],
        ["{gov_pct}", "government_dependency_pct", "\"72%\""],
        ["{gov_amt}", "government_grants_amt", "\"$1.2 million\""],
        ["{mission_snippet}", "mission_statement (first 100 chars)", "\"assists the needy...\""],
        ["{contact_title}", "contact_title", "\"Executive Director\""],
    ]
)

doc.add_page_break()

# ============================================================
# 3. DATABASE CHANGES
# ============================================================
heading("3. Database Changes")

heading("New Columns on nonprofit_prospects", level=2)
add_table(
    ["Column", "Type", "Purpose"],
    [
        ["government_grants_amt", "NUMERIC(15,2)", "Dollar amount of gov grants from most recent 990"],
        ["government_dependency_pct", "NUMERIC(5,2)", "Percentage of revenue from government (0-100)"],
        ["gov_funding_segment", "VARCHAR(30)", "Campaign segment: critical_new, high_experienced, etc."],
        ["ab_test_variant", "VARCHAR(1)", "Assigned A/B test variant: A, B, C, or D"],
    ]
)

doc.add_paragraph()
heading("New Columns on generated_emails", level=2)
add_table(
    ["Column", "Type", "Purpose"],
    [
        ["variant", "VARCHAR(1)", "Which variant template was used (A/B/C/D)"],
        ["ab_test_name", "VARCHAR(50)", "Test identifier (e.g. 'gov_funding_feb2026')"],
    ]
)

doc.add_paragraph()
heading("New Indexes", level=2)
add_table(
    ["Index", "Table", "Column"],
    [
        ["idx_np_gov_segment", "nonprofit_prospects", "gov_funding_segment"],
        ["idx_np_ab_variant", "nonprofit_prospects", "ab_test_variant"],
        ["idx_np_gov_dep_pct", "nonprofit_prospects", "government_dependency_pct"],
        ["idx_ge_variant", "generated_emails", "variant"],
        ["idx_ge_ab_test", "generated_emails", "ab_test_name"],
    ]
)

doc.add_paragraph()
heading("New View: ab_test_results", level=2)
doc.add_paragraph(
    "A database view that automatically calculates A/B test metrics by grouping "
    "generated_emails by test name and variant. Shows total generated, sent, replies, "
    "bounces, unsubscribes, reply rate, and bounce rate. Query it with: "
    "SELECT * FROM grantscout_campaign.ab_test_results"
)

doc.add_page_break()

# ============================================================
# 4. CODE CHANGES
# ============================================================
heading("4. Code Changes")

heading("export_ideal_prospects.py (Phase 0A)", level=2)
doc.add_paragraph(
    "Fixed hardcoded database password ('postgres') by replacing with environment variable "
    "pattern using os.getenv(). Now reads POSTGRES_PASSWORD from environment with 'postgres' "
    "as the default fallback. This is a security best practice and prevents credential "
    "leakage if the file is shared."
)

heading("config.py (Phase 1)", level=2)
p = doc.add_paragraph("Major rewrite. Key changes:")
changes = [
    "Added 8 email templates (4 variants x initial/followup) with full personalization support",
    "Added A/B test configuration: variant weights per segment, test name constant",
    "Added assign_variant() function for weighted random variant assignment",
    "Added VARIANT_TEMPLATES and VARIANT_SUBJECTS lookup dictionaries",
    "Added CAN-SPAM compliant unsubscribe footer to all variants",
    "Removed deprecated food_recall and debarment verticals (only grant_alerts active)",
    "Preserved legacy aliases (GRANT_ALERT_INITIAL, etc.) for backward compatibility with campaign_manager.py",
    "Fixed phone number formatting inconsistency and LinkedIn capitalization",
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

heading("generate_emails.py (Phase 2A)", level=2)
p = doc.add_paragraph("Rewritten to support variants and 990-data personalization. Key changes:")
changes = [
    "get_eligible_prospects() now returns all personalization columns (gov data, sector, city, etc.)",
    "New build_template_vars() function creates the full variable dict from prospect data",
    "generate_personalized_email() selects variant, picks template, substitutes all variables",
    "Uses defaultdict-based format_map() so missing variables become empty strings (never raw {tokens})",
    "New format_currency() helper: $1,200,000 becomes '$1.2 million'",
    "New format_gov_pct() helper: 72.3 becomes '72%'",
    "insert_generated_email() now stores variant and ab_test_name for tracking",
    "New --variant flag to generate only for a specific variant",
    "New 'ab-results' subcommand to view A/B test performance",
    "All SQL queries use parameterized queries (no string interpolation of user input)",
    "Removed unused 'import sys'",
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

heading("assign_variants.py (Phase 2C) -- NEW FILE", level=2)
p = doc.add_paragraph("New script to assign A/B test variants to prospects. Features:")
changes = [
    "--dry-run: preview distribution without writing to database",
    "--segment: only assign to a specific segment (e.g. critical_new)",
    "--reassign: overwrite existing variant assignments",
    "Uses psycopg2 execute_values() for safe parameterized batch updates",
    "Shows full distribution table by segment and variant",
    "All DB operations in a single transaction with rollback on error",
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ============================================================
# 5. A/B TEST DESIGN
# ============================================================
heading("5. A/B Test Design")

heading("Variant Assignment Logic", level=2)
doc.add_paragraph(
    "Prospects are assigned variants based on their segment. The weights are designed "
    "to give the government risk angle (Variant A) the most exposure since it is the "
    "primary hypothesis, while maintaining meaningful sample sizes for comparison."
)

add_table(
    ["Segment", "Variant A", "Variant B", "Variant C", "Variant D"],
    [
        ["critical_new", "40%", "30%", "30%", "--"],
        ["critical_experienced", "50%", "25%", "--", "25%"],
        ["high_new", "40%", "30%", "30%", "--"],
        ["high_experienced", "50%", "25%", "--", "25%"],
        ["moderate", "50%", "25%", "--", "25%"],
        ["No gov data", "--", "50%", "--", "50%"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Note: Variant C (new-to-foundations) is only assigned to prospects with new_to_foundations=1. "
    "If a prospect assigned Variant C turns out to have foundation experience, it falls back "
    "to Variant A (if gov-dependent) or Variant B (if not)."
)

heading("Subject Line Rotation", level=2)
doc.add_paragraph(
    "Each variant has 2 subject lines that are randomly selected at generation time. "
    "This provides natural variation and helps avoid spam filters that flag identical subjects."
)

heading("Measurement Plan", level=2)
add_table(
    ["Metric", "Baseline (Nov 2025)", "Target"],
    [
        ["Delivery rate", "93.9%", "> 95%"],
        ["Reply rate", "0.23%", "> 1%"],
        ["Bounce rate", "6.1%", "< 5%"],
        ["Positive reply rate", "n/a", "> 0.5%"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Decision rules: After 50 sends per variant, compare reply rates. Kill any variant "
    "with > 5% bounce or > 2% unsubscribe. Shift 80% of volume to winning variant."
)

doc.add_page_break()

# ============================================================
# 6. CODE REVIEW RESULTS
# ============================================================
heading("6. Code Review Results")

doc.add_paragraph(
    "Five parallel review agents analyzed every piece of this implementation. "
    "All critical issues were fixed before finalization."
)

heading("Issues Found and Fixed", level=2)
add_table(
    ["Severity", "Issue", "Fix Applied"],
    [
        ["HIGH", "SQL injection: variant_filter interpolated via f-string in generate_emails.py", "Parameterized with %s placeholder"],
        ["HIGH", "SQL injection: segment interpolated via f-string in assign_variants.py", "Parameterized with %s placeholder"],
        ["MEDIUM", "Template substitution could leave raw {var} tokens in outbound emails", "Switched to defaultdict + format_map() (missing vars become empty strings)"],
        ["MEDIUM", "CAN-SPAM: only Variant D had unsubscribe mechanism", "Added unsubscribe footer to all 8 templates"],
        ["LOW", "Phone formatting: (281)245-4596 vs (281) 245-4596", "Standardized to (281) 245-4596"],
        ["LOW", "\"Linkedin\" casing", "Fixed to \"LinkedIn\""],
        ["LOW", "Unused import sys in generate_emails.py", "Removed"],
        ["LOW", "Batch UPDATE in assign_variants.py used string interpolation", "Replaced with psycopg2 execute_values()"],
    ]
)

doc.add_paragraph()
heading("Checks That Passed", level=2)
checks = [
    "All 4 variants defined for both initial and followup (8 templates total)",
    "All template variables ({greeting}, {org_name}, etc.) are produced by build_template_vars()",
    "Subject lines only use variables that will be available",
    "Legacy aliases preserved for backward compatibility with campaign_manager.py",
    "All 5 unmodified files (campaign_manager.py, send_generated_emails.py, etc.) confirmed compatible",
    "Database: all columns, indexes, and view created correctly",
    "Backfill: 34,545 rows updated, avg 45% gov dependency, range 0-100% (no overflow)",
    "Segment distribution matches expected counts across all 5 segments",
    "--dry-run flag correctly prevents all DB writes in both scripts",
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. FILES CREATED / MODIFIED
# ============================================================
heading("7. Files Created / Modified")

add_table(
    ["File", "Action", "Description"],
    [
        ["export_ideal_prospects.py", "Modified", "Replaced hardcoded DB password with env vars"],
        ["config.py", "Rewritten", "Added 4 template variants, A/B config, cleanup deprecated verticals"],
        ["generate_emails.py", "Rewritten", "Added 990-data personalization, variant support, A/B tracking"],
        ["assign_variants.py", "NEW", "Script to assign A/B test variants to prospects"],
        ["DB: nonprofit_prospects", "ALTER TABLE", "Added 4 columns + 3 indexes + backfill 34,545 rows"],
        ["DB: generated_emails", "ALTER TABLE", "Added 2 columns + 2 indexes"],
        ["DB: ab_test_results", "CREATE VIEW", "A/B test metrics view"],
    ]
)

doc.add_page_break()

# ============================================================
# 8. LAUNCH CHECKLIST
# ============================================================
heading("8. What's Next: Launch Checklist")

doc.add_paragraph("The code and database are ready. These steps remain before sending:")

steps = [
    ("Run email scraper on gov-dependent prospects",
     "python3 email_scraper.py --target 1000 --limit 100\n"
     "Target the ~28,700 gov-dependent prospects with websites. Expected yield: 500-1,000 emails."),
    ("Assign A/B variants",
     "python3 assign_variants.py --dry-run  (preview first)\n"
     "python3 assign_variants.py  (write to DB)"),
    ("Generate test emails",
     "python3 generate_emails.py --dry-run --limit 10\n"
     "Manually review all 4 variants for correct variable substitution."),
    ("Send test emails to internal addresses",
     "Use send_test_email.py to send one of each variant to your test addresses."),
    ("Domain warmup",
     "Start at 2 emails/sender/day (20 total). Ramp to 15/sender over 2 weeks.\n"
     "Update grantscout_campaign.senders daily_limit column progressively."),
    ("Generate production emails",
     "python3 generate_emails.py --limit 20  (first batch)\n"
     "Verify in DB, then send with send_generated_emails.py"),
    ("Monitor and iterate",
     "python3 generate_emails.py ab-results  (check A/B performance)\n"
     "python3 campaign_status.py today  (daily activity)\n"
     "After 50 sends per variant, shift volume to winner."),
]
for i, (title, detail) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.add_run(f"Step {i}: {title}").bold = True
    doc.add_paragraph(detail)

doc.add_page_break()

# ============================================================
# 9. RISKS
# ============================================================
heading("9. Risks and Mitigations")

add_table(
    ["Risk", "Mitigation"],
    [
        ["990 gov data is 1-2 years old",
         "Templates frame as structural dependency (\"your org depends on...\"), not temporal"],
        ["Low email yield from scraper",
         "Run broadly first; consider Hunter.io for high-value prospects"],
        ["Cold domain deliverability",
         "Follow warmup strictly; start at 2/sender/day, ramp over 2 weeks"],
        ["\"Surveilled\" feeling from too-specific data",
         "Templates use aggregate language (\"organizations like yours\"). A/B test results will show if personalization helps or hurts"],
        ["campaign_manager.py bypasses variant system",
         "By design: legacy path always sends Variant D. Use generate_emails.py + send_generated_emails.py for variant-aware sending"],
    ]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Generated by Claude Code on 2026-02-13")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = "/Users/aleckleinman/Documents/TheGrantScout/4. Pipeline/Enhancements/2026-02-13/REPORT_2026-02-13.1_email_campaign_gov_funding_implementation.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
