#!/usr/bin/env python3
"""Generate Word doc for bmf column reference with full code tables."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

TOTAL = 1935635
HEADER_BG = '1e3a5f'


def add_styled_header_row(table, headers):
    """Apply navy background + white bold text to the first row of a table."""
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): HEADER_BG,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)


def add_code_table(doc, title, rows, col_headers=None):
    """Add a code lookup table with title. rows = list of tuples (code, count, meaning)."""
    if col_headers is None:
        col_headers = ['Code', 'Count', 'Meaning']
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    table = doc.add_table(rows=1 + len(rows), cols=len(col_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_styled_header_row(table, col_headers)
    for row_idx, vals in enumerate(rows):
        for col_idx, val in enumerate(vals):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


# ── Column reference data ──────────────────────────────────────────────

sections_data = [
    ("Organization Identity", [
        ("ein", "VARCHAR(20)", 1935635, "Employer Identification Number (9-digit, no dashes). Primary key."),
        ("name", "TEXT", 1935635, "Legal name of exempt organization"),
        ("sort_name", "TEXT", 415722, "Secondary/sort name (often inverted: 'Red Cross, American')"),
        ("ico", "TEXT", 1242531, "In Care Of name (contact person at org address)"),
    ]),
    ("Address", [
        ("street", "TEXT", 1935635, "Street address or PO Box"),
        ("city", "TEXT", 1935635, "City"),
        ("state", "VARCHAR(2)", 1934272, "State code (2-letter). Includes DC, PR, territories."),
        ("zip", "VARCHAR(10)", 1935635, "ZIP code (5-digit or ZIP+4)"),
    ]),
    ("Tax Exemption Classification", [
        ("subsection", "VARCHAR(5)", 1935635, "IRC subsection code (see Subsection Codes below)"),
        ("foundation", "VARCHAR(5)", 1935635, "Foundation status code (see Foundation Codes below)"),
        ("deductibility", "VARCHAR(5)", 1935635, "Contribution deductibility (see Deductibility Codes below)"),
        ("status", "VARCHAR(5)", 1935635, "Exempt organization status (see Status Codes below)"),
        ("ruling", "VARCHAR(6)", 1935635, "IRS ruling/determination date (YYYYMM format). 1.92M have a meaningful date (not 000000)."),
    ]),
    ("Organization Structure", [
        ("organization", "VARCHAR(5)", 1935635, "Organization type (see Organization Codes below)"),
        ("affiliation", "VARCHAR(5)", 1935635, "Affiliation relationship (see Affiliation Codes below)"),
        ("group_code", "VARCHAR(10)", 1935635, "Group exemption number. 0000 = no group (1.53M orgs). Non-zero = part of a group exemption (e.g., local chapters of a national org)."),
        ("classification", "VARCHAR(10)", 1935635, "IRS classification codes (legacy). Multi-digit code describing the specific exemption basis."),
        ("activity", "VARCHAR(10)", 1935635, "Activity codes (legacy, superseded by NTEE). Up to 3 three-digit codes describing what the org does. 715K have meaningful values."),
        ("ntee_cd", "VARCHAR(10)", 1352945, "National Taxonomy of Exempt Entities code (see NTEE Codes below)"),
    ]),
    ("Financials", [
        ("asset_amt", "NUMERIC(15,2)", 1482390, "Total assets from the organization's most recent Form 990 filing"),
        ("income_amt", "NUMERIC(15,2)", 1482390, "Gross income from the most recent Form 990 filing"),
        ("revenue_amt", "NUMERIC(15,2)", 1356504, "Total revenue (Form 990 Line 12 or 990-EZ Line 9). The most reliable revenue figure in the BMF."),
        ("asset_cd", "VARCHAR(5)", 1935635, "Asset size code (see Asset/Income Codes below)"),
        ("income_cd", "VARCHAR(5)", 1935635, "Income size code (same ranges as asset_cd)"),
    ]),
    ("Filing Information", [
        ("tax_period", "VARCHAR(6)", 1502230, "Most recent tax period end date (YYYYMM). A recent date means the org is actively filing."),
        ("filing_req_cd", "VARCHAR(5)", 1935635, "Form 990 filing requirement (see Filing Requirement Codes below)"),
        ("pf_filing_req_cd", "VARCHAR(5)", 1935635, "Private foundation filing requirement (see PF Filing Codes below)"),
        ("acct_pd", "VARCHAR(5)", 1935635, "Accounting period end month (01-12). 12 = calendar year filer."),
    ]),
    ("Metadata", [
        ("source_file", "VARCHAR(10)", 1935635, "Which IRS EO BMF regional file this came from (eo1, eo2, eo3, or eo4)"),
        ("created_at", "TIMESTAMP", 1935635, "When this row was inserted into our database"),
    ]),
]

# ── Build document ─────────────────────────────────────────────────────

# Title
title = doc.add_heading('bmf - Column Reference', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

# Subtitle info
p = doc.add_paragraph()
p.add_run('Table: ').bold = True
p.add_run('f990_2025.bmf\n')
p.add_run('Granularity: ').bold = True
p.add_run('One row per exempt organization (EIN is primary key)\n')
p.add_run('Total rows: ').bold = True
p.add_run('1,935,635\n')
p.add_run('Source: ').bold = True
p.add_run('IRS Exempt Organizations Business Master File (EO BMF), 4 regional files\n')
p.add_run('As of: ').bold = True
p.add_run('February 2026')

# ── Column tables ──────────────────────────────────────────────────────

for section_name, rows in sections_data:
    doc.add_heading(section_name, level=2)

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_styled_header_row(table, ['Column', 'Type', 'Populated', '%', 'Meaning'])

    for row_idx, (col_name, col_type, count, meaning) in enumerate(rows):
        pct = f"{count / TOTAL * 100:.1f}%" if count > 0 else "0%"
        values = [col_name, col_type, f"{count:,}", pct, meaning]
        for col_idx, val in enumerate(values):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if count < TOTAL * 0.25:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(0.9)
        row.cells[3].width = Inches(0.5)
        row.cells[4].width = Inches(3.3)

    doc.add_paragraph()

# ── Code Lookup Tables ─────────────────────────────────────────────────

doc.add_heading('Code Lookup Tables', level=1)
h = doc.paragraphs[-1]
h.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

# Foundation codes
add_code_table(doc, 'Foundation Codes (foundation column)', [
    ('00', '325,331', 'Not a 501(c)(3) organization (field not applicable)'),
    ('02', '310', 'Private operating foundation exempt from excise taxes on net investment income'),
    ('03', '9,086', 'Private operating foundation (all other)'),
    ('04', '124,802', 'Private non-operating foundation (the classic grantmaking PF)'),
    ('10', '290,055', 'Church (170(b)(1)(A)(i))'),
    ('11', '24,124', 'School (170(b)(1)(A)(ii))'),
    ('12', '5,808', 'Hospital or cooperative hospital service org (170(b)(1)(A)(iii))'),
    ('13', '2,645', 'Org operated for benefit of a college or university (170(b)(1)(A)(iv))'),
    ('14', '361', 'Governmental unit (170(b)(1)(A)(v))'),
    ('15', '658,420', 'Publicly supported org receiving >1/3 from contributions (509(a)(1)). Example: United Way, Salvation Army.'),
    ('16', '470,916', 'Publicly supported org receiving >1/3 from gross receipts or operated for benefit of 509(a)(1) orgs (509(a)(2)). Example: YMCA, museum gift shops.'),
    ('17', '15,380', 'Organization determined publicly supported by IRS (alternative test)'),
    ('21', '5,470', 'Supporting organization Type I (509(a)(3)). Operated/supervised by supported org.'),
    ('22', '2,073', 'Supporting organization Type II (509(a)(3)). Supervised or controlled in connection with.'),
    ('23', '1,318', 'Supporting organization Type III functionally integrated (509(a)(3))'),
    ('24', '704', 'Supporting organization Type III other (509(a)(3))'),
])

# Subsection codes
add_code_table(doc, 'Subsection Codes (subsection column)', [
    ('03', '1,610,053', '501(c)(3) - Charitable, educational, religious, scientific. The main nonprofit category. Eligible for tax-deductible donations.'),
    ('04', '69,191', '501(c)(4) - Civic leagues, social welfare orgs. Example: Sierra Club, NRA. Can lobby.'),
    ('05', '43,472', '501(c)(5) - Labor, agricultural, horticultural orgs. Example: AFL-CIO locals, farm bureaus.'),
    ('06', '57,647', '501(c)(6) - Business leagues, chambers of commerce, trade associations. Example: local Chamber of Commerce.'),
    ('07', '46,686', '501(c)(7) - Social and recreational clubs. Example: country clubs, yacht clubs.'),
    ('08', '34,956', '501(c)(8) - Fraternal beneficiary societies. Example: Knights of Columbus, Elks lodges.'),
    ('09', '5,268', '501(c)(9) - Voluntary employees beneficiary associations (VEBAs). Employee benefit plans.'),
    ('10', '14,365', '501(c)(10) - Domestic fraternal societies. Similar to (8) but organized under lodge system.'),
    ('12', '5,401', '501(c)(12) - Benevolent life insurance associations, mutual telephone/electric companies.'),
    ('13', '9,688', '501(c)(13) - Cemetery companies. Nonprofit cemeteries.'),
    ('14', '1,400', '501(c)(14) - Credit unions, mutual reserve funds.'),
    ('19', '24,740', '501(c)(19) - Veterans organizations. Example: VFW, American Legion posts.'),
    ('25', '521', '501(c)(25) - Title holding corps for multiple exempt orgs.'),
    ('92', '5,830', '527 - Political organizations. Campaign committees, PACs.'),
])

# Status codes
add_code_table(doc, 'Status Codes (status column)', [
    ('01', '1,927,507', 'Unconditional Exemption. Active and in good standing with IRS.'),
    ('02', '644', 'Conditional Exemption. Exempt with conditions or limitations.'),
    ('12', '6,645', 'Revoked. Exemption revoked, usually for failure to file for 3 consecutive years (auto-revocation).'),
    ('25', '839', 'Terminated. Voluntarily gave up exempt status or IRS terminated it.'),
])

# Deductibility codes
add_code_table(doc, 'Deductibility Codes (deductibility column)', [
    ('1', '1,666,025', 'Contributions are deductible. Standard charitable deduction. Example: donations to a 501(c)(3) public charity.'),
    ('2', '235,361', 'Contributions are deductible by treaty or election. Limited deductibility. Example: certain foreign orgs.'),
    ('0', '33,716', 'Contributions are NOT deductible. Example: social clubs (501(c)(7)), business leagues (501(c)(6)).'),
    ('4', '533', 'Deductibility depends on date of gift. Applies to orgs with unusual deduction provisions.'),
])

# Organization type codes
add_code_table(doc, 'Organization Type Codes (organization column)', [
    ('1', '1,465,391', 'Corporation. The most common legal form for nonprofits. Example: most 501(c)(3) charities.'),
    ('2', '48,443', 'Trust. Common for private foundations. Example: The Ford Foundation Trust.'),
    ('3', '1,540', 'Co-operative. Member-owned organizations.'),
    ('4', '102', 'Partnership. Rare for exempt orgs.'),
    ('5', '407,089', 'Association. Unincorporated membership organizations. Example: many churches, VFW posts.'),
    ('6', '2,679', 'Other. Does not fit standard categories.'),
    ('0', '10,391', 'Not specified or unknown.'),
])

# Affiliation codes
add_code_table(doc, 'Affiliation Codes (affiliation column)', [
    ('3', '1,522,995', 'Independent. Not part of a group exemption. Files its own Form 990. The majority of orgs.'),
    ('9', '398,540', 'Subordinate. Part of a group exemption held by a central org. Example: local Girl Scout council under national org.'),
    ('1', '3,468', 'Central organization (group ruling). The parent that holds the group exemption. Example: Girl Scouts of the USA national.'),
    ('6', '3,279', 'Central organization that is NOT included in its own group exemption group return.'),
    ('0', '5,512', 'Not specified.'),
    ('2', '1,042', 'Intermediate organization (group return). Files a group return that includes subordinates.'),
    ('8', '715', 'Central organization whose group exemption has been revoked.'),
    ('7', '84', 'Intermediate organization (not filing group return).'),
])

# Asset/Income size codes
add_code_table(doc, 'Asset Size Codes (asset_cd column) and Income Size Codes (income_cd column)', [
    ('0', '1,221,763 / 1,232,066', '$0 or not reported. Most small orgs and churches that do not file Form 990.'),
    ('1', '75,763 / 49,799', '$1 to $10,000'),
    ('2', '42,590 / 31,093', '$10,000 to $25,000'),
    ('3', '121,327 / 149,410', '$25,000 to $100,000'),
    ('4', '184,922 / 237,532', '$100,000 to $500,000'),
    ('5', '74,472 / 71,227', '$500,000 to $1,000,000 ($1M)'),
    ('6', '121,973 / 99,511', '$1,000,000 to $5,000,000 ($5M)'),
    ('7', '32,067 / 24,106', '$5,000,000 to $10,000,000 ($10M)'),
    ('8', '41,372 / 29,060', '$10,000,000 to $50,000,000 ($50M)'),
    ('9', '19,386 / 11,831', '$50,000,000+ ($50M and above)'),
], col_headers=['Code', 'Count (Asset / Income)', 'Range'])

# Filing requirement codes
add_code_table(doc, 'Filing Requirement Codes (filing_req_cd column)', [
    ('01', '505,307', 'Form 990 required. Gross receipts >= $200K or assets >= $500K.'),
    ('02', '986,096', 'Form 990 or 990-EZ. Gross receipts < $200K and assets < $500K. Can file the shorter 990-EZ.'),
    ('03', '813', 'Form 990-PF required. Private foundations (rare in this field, most PFs use pf_filing_req_cd).'),
    ('06', '284,822', 'Not required to file. Churches, government agencies, and certain religious orgs are exempt from filing.'),
    ('00', '146,018', 'Form 990-N (e-Postcard) eligible. Gross receipts normally <= $50K. Very small orgs.'),
    ('07', '766', 'Section 4947(a)(1) nonexempt charitable trust treated as private foundation.'),
    ('13', '9,153', 'Section 4947(a)(1) nonexempt charitable trust NOT treated as private foundation.'),
    ('14', '2,659', 'Not required to file (other exemptions from filing).'),
    ('04', '1', 'Form 990-BL required (very rare).'),
])

# PF filing requirement codes
add_code_table(doc, 'PF Filing Requirement Codes (pf_filing_req_cd column)', [
    ('0', '1,796,161', 'Not required to file a 990-PF. The vast majority, as only private foundations file 990-PFs.'),
    ('1', '139,260', 'Form 990-PF required. This is how to identify private foundations in the BMF.'),
    ('3', '191', 'Section 4947(a)(1) nonexempt charitable trust filing 990-PF.'),
    ('2', '23', 'Form 990-PF group return.'),
])

# NTEE codes overview
add_code_table(doc, 'NTEE Major Categories (first letter of ntee_cd column)', [
    ('A', '', 'Arts, Culture, Humanities. Example: A20 = Museums, A60 = Performing Arts.'),
    ('B', '', 'Education. Example: B20 = Elementary/Secondary Schools, B40 = Higher Education.'),
    ('C', '', 'Environment. Example: C20 = Pollution Abatement, C30 = Natural Resources.'),
    ('D', '', 'Animal-Related. Example: D20 = Animal Protection, D30 = Wildlife.'),
    ('E', '', 'Health Care. Example: E20 = Hospitals, E40 = Reproductive Health.'),
    ('F', '', 'Mental Health & Crisis. Example: F20 = Substance Abuse, F30 = Mental Health.'),
    ('G', '', 'Diseases & Disorders. Example: G20 = Birth Defects, G40 = Cancer.'),
    ('H', '', 'Medical Research. Example: H20 = Birth Defects Research.'),
    ('I', '', 'Crime & Legal. Example: I20 = Crime Prevention, I40 = Rehabilitation.'),
    ('J', '', 'Employment. Example: J20 = Job Training, J30 = Vocational Rehabilitation.'),
    ('K', '', 'Food, Agriculture, Nutrition. Example: K20 = Food Banks, K30 = Agriculture.'),
    ('L', '', 'Housing & Shelter. Example: L20 = Housing Development, L40 = Homeless Services.'),
    ('M', '', 'Public Safety & Disaster. Example: M20 = Disaster Preparedness, M40 = Search/Rescue.'),
    ('N', '', 'Recreation & Sports. Example: N20 = Camps, N60 = Amateur Sports.'),
    ('O', '', 'Youth Development. Example: O20 = Youth Centers, O40 = Scouting.'),
    ('P', '', 'Human Services (Multipurpose). Example: P20 = Multipurpose Service, P80 = Emergency Assistance.'),
    ('Q', '', 'International Affairs. Example: Q20 = Promotion of International Understanding.'),
    ('R', '', 'Civil Rights & Social Action. Example: R20 = Intergroup Relations, R60 = Civil Liberties.'),
    ('S', '', 'Community Improvement. Example: S20 = Community Coalitions, S40 = Business & Industry.'),
    ('T', '', 'Philanthropy & Grantmaking. Example: T20 = Private Grantmaking Foundations, T30 = Public Foundations.'),
    ('U', '', 'Science & Technology. Example: U20 = General Science, U40 = Engineering.'),
    ('V', '', 'Social Science. Example: V20 = Social Science Research.'),
    ('W', '', 'Public/Society Benefit. Example: W20 = Government & Public Admin.'),
    ('X', '', 'Religion-Related. Example: X20 = Christian, X30 = Jewish.'),
    ('Y', '', 'Mutual & Membership Benefit. Example: Y20 = Insurance, Y40 = Pensions.'),
    ('Z', '', 'Unknown / Unclassified.'),
], col_headers=['Letter', '', 'Category and Examples'])

# ── Summary Stats ──────────────────────────────────────────────────────

doc.add_heading('Summary Stats', level=2)
stats_data = [
    ('Total organizations', '1,935,635'),
    ('501(c)(3) orgs', '1,610,053 (83.2%)'),
    ('  Public charities (15+16+17)', '1,144,716 (59.1%)'),
    ('  Churches (10)', '290,055 (15.0%)'),
    ('  Private foundations (02+03+04)', '134,198 (6.9%)'),
    ('  Schools (11)', '24,124 (1.2%)'),
    ('  Supporting orgs (21-24)', '9,565 (0.5%)'),
    ('  Hospitals (12)', '5,808 (0.3%)'),
    ('Non-501(c)(3) orgs', '325,582 (16.8%)'),
    ('Active (status 01)', '1,927,507 (99.6%)'),
    ('With NTEE code', '1,352,945 (69.9%)'),
    ('With financial data (asset_amt)', '1,482,390 (76.6%)'),
    ('Unique states/territories', '62'),
]
stats_table = doc.add_table(rows=1 + len(stats_data), cols=2)
stats_table.style = 'Table Grid'
add_styled_header_row(stats_table, ['Metric', 'Value'])

for row_idx, (metric, value) in enumerate(stats_data):
    stats_table.rows[row_idx + 1].cells[0].text = metric
    stats_table.rows[row_idx + 1].cells[1].text = value
    for cell in stats_table.rows[row_idx + 1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

output_path = '/Users/aleckleinman/Documents/TheGrantScout/2. Docs/bmf_column_reference.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
