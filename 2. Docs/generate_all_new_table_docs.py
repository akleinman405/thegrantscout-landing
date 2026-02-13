#!/usr/bin/env python3
"""
Generate column reference docs (.md + .docx) for 14 new import tables.
Queries PostgreSQL for column metadata and population rates.
Matches format of existing pf_returns_column_reference.md/.docx.
"""

import psycopg2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DB_CONFIG = {
    'host': 'localhost',
    'port': 5432,
    'database': 'thegrantscout',
    'user': 'postgres',
    'password': 'kmalec21',
}

OUTPUT_DIR = '/Users/aleckleinman/Documents/TheGrantScout/2. Docs'

HEADER_BG = '1e3a5f'

# ── Table definitions: name, granularity description, section groupings ──

TABLE_DEFS = {
    'schedule_i_grants': {
        'granularity': 'One row per grant from 990 Schedule I (grants to domestic orgs/individuals)',
        'source': '990/990-EZ Schedule I XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'filer_ein', 'tax_year', 'created_at'],
            'Recipient Info': ['recipient_ein', 'recipient_name', 'recipient_city', 'recipient_state', 'recipient_zip'],
            'Grant Details': ['irc_section', 'amount', 'noncash_amount', 'purpose', 'grant_type'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'filer_ein': 'EIN of the granting organization',
            'tax_year': 'Tax year of the filing',
            'recipient_ein': 'EIN of the grant recipient (if available)',
            'recipient_name': 'Name of the grant recipient',
            'recipient_city': 'City of the grant recipient',
            'recipient_state': 'State code of the grant recipient (2-letter)',
            'recipient_zip': 'ZIP code of the grant recipient',
            'irc_section': 'IRC section under which recipient is exempt (e.g., 501(c)(3))',
            'amount': 'Cash grant amount in dollars',
            'noncash_amount': 'Non-cash assistance amount in dollars',
            'purpose': 'Purpose or description of the grant',
            'grant_type': 'Type of grant (e.g., cash, non-cash)',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_a': {
        'granularity': 'One row per 990 filing (public charity status and support test)',
        'source': '990 Schedule A XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Organization Type': ['church_ind', 'school_ind', 'hospital_ind', 'medical_research_org_ind',
                                  'college_org_ind', 'governmental_unit_ind', 'public_safety_org_ind',
                                  'community_trust_ind'],
            'Public Charity Status': ['public_charity_509a1_ind', 'public_charity_509a2_ind',
                                       'supporting_org_509a3_ind', 'supporting_org_type1_ind',
                                       'supporting_org_type2_ind', 'supporting_org_type3_func_int_ind',
                                       'supporting_org_type3_non_func_int_ind'],
            'Public Support Test': ['public_support_cy170_pct', 'public_support_total_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'church_ind': 'Organization is a church (170(b)(1)(A)(i))',
            'school_ind': 'Organization is a school (170(b)(1)(A)(ii))',
            'hospital_ind': 'Organization is a hospital (170(b)(1)(A)(iii))',
            'medical_research_org_ind': 'Medical research org (170(b)(1)(A)(iii))',
            'college_org_ind': 'College/university support org (170(b)(1)(A)(iv))',
            'governmental_unit_ind': 'Governmental unit (170(b)(1)(A)(v))',
            'public_safety_org_ind': 'Public safety testing org (509(a)(4))',
            'community_trust_ind': 'Community trust (170(b)(1)(A)(vi))',
            'public_charity_509a1_ind': 'Publicly supported charity under 509(a)(1)',
            'public_charity_509a2_ind': 'Publicly supported charity under 509(a)(2)',
            'supporting_org_509a3_ind': 'Supporting organization under 509(a)(3)',
            'supporting_org_type1_ind': 'Supporting org Type I (operated/supervised by)',
            'supporting_org_type2_ind': 'Supporting org Type II (supervised/controlled with)',
            'supporting_org_type3_func_int_ind': 'Supporting org Type III functionally integrated',
            'supporting_org_type3_non_func_int_ind': 'Supporting org Type III non-functionally integrated',
            'public_support_cy170_pct': 'Public support percentage (170 test, current year)',
            'public_support_total_amt': 'Total public support amount',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'pf_capital_gains': {
        'granularity': 'One row per capital gain/loss line item from 990-PF Part IV',
        'source': '990-PF Part IV XML filings',
        'sections': {
            'Record Metadata': ['id', 'pf_return_id', 'ein', 'tax_year', 'created_at'],
            'Property Details': ['property_desc', 'how_acquired_cd', 'date_acquired', 'date_sold'],
            'Financial': ['gross_sales_price_amt', 'cost_or_other_basis_amt', 'gain_or_loss_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'pf_return_id': 'FK to pf_returns.id',
            'ein': 'Foundation EIN',
            'tax_year': 'Tax year of the filing',
            'property_desc': 'Description of the property sold (e.g., stock name, real estate)',
            'how_acquired_cd': 'How property was acquired: P=Purchase, D=Donation',
            'date_acquired': 'Date property was acquired',
            'date_sold': 'Date property was sold',
            'gross_sales_price_amt': 'Gross sales price of the property',
            'cost_or_other_basis_amt': 'Cost or other basis of the property',
            'gain_or_loss_amt': 'Net capital gain or loss (sales price minus basis)',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_o_narratives': {
        'granularity': 'One row per supplemental explanation from Schedule O',
        'source': '990/990-EZ Schedule O XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Narrative Content': ['form_and_line_ref', 'explanation_txt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'form_and_line_ref': 'Form and line number this explanation refers to (e.g., "Form 990, Part III, Line 4d")',
            'explanation_txt': 'Full text of the supplemental explanation',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'np_balance_sheet': {
        'granularity': 'One row per balance sheet line item from 990/990-EZ Part X',
        'source': '990/990-EZ Part X XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Balance Sheet Data': ['line_item', 'boy_amt', 'eoy_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'line_item': 'Balance sheet line item code (e.g., cash, investments, payables)',
            'boy_amt': 'Beginning-of-year amount',
            'eoy_amt': 'End-of-year amount',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'pf_investments': {
        'granularity': 'One row per investment holding from 990-PF Part II',
        'source': '990-PF Part II XML filings',
        'sections': {
            'Record Metadata': ['id', 'pf_return_id', 'ein', 'tax_year', 'created_at'],
            'Investment Details': ['investment_type', 'description', 'book_value_eoy_amt', 'fmv_eoy_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'pf_return_id': 'FK to pf_returns.id',
            'ein': 'Foundation EIN',
            'tax_year': 'Tax year of the filing',
            'investment_type': 'Category: corporate_stock, corporate_bond, govt_obligations, land_buildings, etc.',
            'description': 'Description of the specific investment holding',
            'book_value_eoy_amt': 'Book value at end of year',
            'fmv_eoy_amt': 'Fair market value at end of year',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'pf_revenue_expenses': {
        'granularity': 'One row per revenue/expense line item from 990-PF Part I',
        'source': '990-PF Part I XML filings',
        'sections': {
            'Record Metadata': ['id', 'pf_return_id', 'ein', 'tax_year', 'created_at'],
            'Line Item Data': ['line_item', 'rev_and_expenses_amt', 'net_invst_incm_amt',
                               'adj_net_incm_amt', 'dsbrs_chrtbl_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'pf_return_id': 'FK to pf_returns.id',
            'ein': 'Foundation EIN',
            'tax_year': 'Tax year of the filing',
            'line_item': 'Revenue/expense line item code (e.g., contributions, dividends, comp_officers)',
            'rev_and_expenses_amt': 'Revenue and expenses per books amount (Column A)',
            'net_invst_incm_amt': 'Net investment income amount (Column B)',
            'adj_net_incm_amt': 'Adjusted net income amount (Column C)',
            'dsbrs_chrtbl_amt': 'Disbursements for charitable purposes amount (Column D)',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_r_related_orgs': {
        'granularity': 'One row per related organization from Schedule R',
        'source': '990 Schedule R XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'filer_ein', 'tax_year', 'created_at'],
            'Related Org Identity': ['schedule_part', 'related_org_name', 'related_org_ein', 'related_org_state'],
            'Relationship Details': ['exempt_code_section', 'public_charity_status',
                                     'direct_controlling_name', 'relationship_type'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'filer_ein': 'EIN of the filing organization',
            'tax_year': 'Tax year of the filing',
            'schedule_part': 'Which part of Schedule R (II, III, or IV)',
            'related_org_name': 'Name of the related organization',
            'related_org_ein': 'EIN of the related organization',
            'related_org_state': 'State code of the related organization',
            'exempt_code_section': 'IRC exempt code section of related org',
            'public_charity_status': 'Public charity status of related org',
            'direct_controlling_name': 'Name of directly controlling entity',
            'relationship_type': 'Type of relationship (e.g., parent, subsidiary)',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'pf_balance_sheet': {
        'granularity': 'One row per balance sheet line item from 990-PF Part II',
        'source': '990-PF Part II XML filings',
        'sections': {
            'Record Metadata': ['id', 'pf_return_id', 'ein', 'tax_year', 'created_at'],
            'Balance Sheet Data': ['line_item', 'boy_amt', 'eoy_amt', 'eoy_fmv_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'pf_return_id': 'FK to pf_returns.id',
            'ein': 'Foundation EIN',
            'tax_year': 'Tax year of the filing',
            'line_item': 'Balance sheet line item code (e.g., cash, investments, payables)',
            'boy_amt': 'Beginning-of-year book value amount',
            'eoy_amt': 'End-of-year book value amount',
            'eoy_fmv_amt': 'End-of-year fair market value amount',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'np_functional_expenses': {
        'granularity': 'One row per filing with functional expense breakdowns (990 Part IX)',
        'source': '990 Part IX XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Grants': ['grants_to_domestic_orgs_amt', 'grants_to_domestic_individuals_amt',
                       'grants_to_foreign_orgs_amt'],
            'Compensation': ['compensation_current_officers_amt', 'compensation_disqualified_amt',
                             'other_salaries_wages_amt', 'pension_plan_contributions_amt',
                             'other_employee_benefits_amt', 'payroll_taxes_amt'],
            'Professional Fees': ['management_fees_amt', 'legal_fees_amt', 'accounting_fees_amt',
                                  'lobbying_fees_amt', 'fundraising_fees_amt', 'investment_mgmt_fees_amt'],
            'Operations': ['advertising_promotion_amt', 'office_expenses_amt', 'info_technology_amt',
                           'occupancy_amt', 'travel_amt', 'conferences_amt', 'interest_amt',
                           'depreciation_amt', 'insurance_amt'],
            'Totals': ['total_functional_expenses_amt', 'total_joint_costs_amt',
                       'total_program_services_amt', 'total_management_general_amt',
                       'total_fundraising_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'grants_to_domestic_orgs_amt': 'Grants to domestic organizations',
            'grants_to_domestic_individuals_amt': 'Grants to domestic individuals',
            'grants_to_foreign_orgs_amt': 'Grants to foreign organizations',
            'compensation_current_officers_amt': 'Compensation of current officers/directors/trustees',
            'compensation_disqualified_amt': 'Compensation to disqualified persons',
            'other_salaries_wages_amt': 'Other salaries and wages',
            'pension_plan_contributions_amt': 'Pension plan accruals and contributions',
            'other_employee_benefits_amt': 'Other employee benefits',
            'payroll_taxes_amt': 'Payroll taxes',
            'management_fees_amt': 'Management and general fees',
            'legal_fees_amt': 'Legal fees',
            'accounting_fees_amt': 'Accounting fees',
            'lobbying_fees_amt': 'Lobbying fees',
            'fundraising_fees_amt': 'Professional fundraising fees',
            'investment_mgmt_fees_amt': 'Investment management fees',
            'advertising_promotion_amt': 'Advertising and promotion',
            'office_expenses_amt': 'Office expenses',
            'info_technology_amt': 'Information technology',
            'occupancy_amt': 'Occupancy expenses',
            'travel_amt': 'Travel',
            'conferences_amt': 'Conferences, conventions, and meetings',
            'interest_amt': 'Interest expense',
            'depreciation_amt': 'Depreciation, depletion, and amortization',
            'insurance_amt': 'Insurance',
            'total_functional_expenses_amt': 'Total functional expenses (all categories)',
            'total_joint_costs_amt': 'Total joint costs (allocated across functions)',
            'total_program_services_amt': 'Total program services expenses',
            'total_management_general_amt': 'Total management and general expenses',
            'total_fundraising_amt': 'Total fundraising expenses',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_j_compensation': {
        'granularity': 'One row per highly-compensated person from Schedule J',
        'source': '990 Schedule J Part II XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Person Info': ['person_nm', 'title_txt'],
            'Compensation Breakdown': ['base_compensation_org_amt', 'bonus_filing_org_amt',
                                        'other_compensation_org_amt', 'deferred_comp_org_amt',
                                        'nontaxable_benefits_amt', 'total_compensation_org_amt',
                                        'comp_reported_related_org_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'person_nm': 'Name of the compensated person',
            'title_txt': 'Title/position of the compensated person',
            'base_compensation_org_amt': 'Base compensation from the filing organization',
            'bonus_filing_org_amt': 'Bonus and incentive compensation from filing org',
            'other_compensation_org_amt': 'Other reportable compensation from filing org',
            'deferred_comp_org_amt': 'Deferred compensation from filing org',
            'nontaxable_benefits_amt': 'Nontaxable benefits (e.g., health insurance, housing)',
            'total_compensation_org_amt': 'Total compensation from the filing organization',
            'comp_reported_related_org_amt': 'Compensation reported by related organizations',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'contractors': {
        'granularity': 'One row per top contractor (highest-compensated, from 990/990-PF)',
        'source': '990 Part VII Section B / 990-PF Part VIII XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'pf_return_id', 'ein', 'tax_year', 'created_at'],
            'Contractor Info': ['contractor_name', 'contractor_address', 'contractor_state'],
            'Service Details': ['service_desc', 'compensation_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id (NULL if from 990-PF)',
            'pf_return_id': 'FK to pf_returns.id (NULL if from 990/990-EZ)',
            'ein': 'EIN of the filing organization',
            'tax_year': 'Tax year of the filing',
            'contractor_name': 'Name of the contractor',
            'contractor_address': 'Address of the contractor',
            'contractor_state': 'State code of the contractor (2-letter)',
            'service_desc': 'Description of services provided',
            'compensation_amt': 'Total compensation paid to the contractor',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_f_grants': {
        'granularity': 'One row per foreign grant region from Schedule F Part I',
        'source': '990 Schedule F Part I XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'filer_ein', 'tax_year', 'created_at'],
            'Region Info': ['region', 'recipient_cnt'],
            'Grant Amounts': ['cash_grant_amt', 'noncash_amt'],
            'Activity Details': ['manner_of_cash', 'activity_desc'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'filer_ein': 'EIN of the granting organization',
            'tax_year': 'Tax year of the filing',
            'region': 'Geographic region of foreign activity (e.g., "Sub-Saharan Africa")',
            'recipient_cnt': 'Number of recipients in this region',
            'cash_grant_amt': 'Total cash grants to this region',
            'noncash_amt': 'Total non-cash assistance to this region',
            'manner_of_cash': 'How cash was disbursed (e.g., wire transfer, check)',
            'activity_desc': 'Description of activities conducted in this region',
            'created_at': 'When this row was inserted into our DB',
        },
    },
    'schedule_d_endowments': {
        'granularity': "One row per organization's endowment data from Schedule D Part V",
        'source': '990 Schedule D Part V XML filings',
        'sections': {
            'Record Metadata': ['id', 'np_return_id', 'ein', 'tax_year', 'created_at'],
            'Current Year Endowment': ['endowment_boy_amt', 'contributions_amt', 'investment_earnings_amt',
                                        'grants_or_scholarships_amt', 'other_expenditures_amt',
                                        'admin_expenses_amt', 'endowment_eoy_amt'],
            'Endowment Composition': ['board_designated_pct', 'perm_restricted_pct', 'temp_restricted_pct'],
            'Year -1 Endowment': ['yr1_endowment_boy_amt', 'yr1_contributions_amt',
                                   'yr1_investment_earnings_amt', 'yr1_grants_or_scholarships_amt',
                                   'yr1_other_expenditures_amt', 'yr1_admin_expenses_amt',
                                   'yr1_endowment_eoy_amt'],
            'Year -2 Endowment': ['yr2_endowment_boy_amt', 'yr2_contributions_amt',
                                   'yr2_investment_earnings_amt', 'yr2_grants_or_scholarships_amt',
                                   'yr2_other_expenditures_amt', 'yr2_admin_expenses_amt',
                                   'yr2_endowment_eoy_amt'],
            'Year -3 Endowment': ['yr3_endowment_boy_amt', 'yr3_contributions_amt',
                                   'yr3_investment_earnings_amt', 'yr3_grants_or_scholarships_amt',
                                   'yr3_other_expenditures_amt', 'yr3_admin_expenses_amt',
                                   'yr3_endowment_eoy_amt'],
            'Year -4 Endowment': ['yr4_endowment_boy_amt', 'yr4_contributions_amt',
                                   'yr4_investment_earnings_amt', 'yr4_grants_or_scholarships_amt',
                                   'yr4_other_expenditures_amt', 'yr4_admin_expenses_amt',
                                   'yr4_endowment_eoy_amt'],
            'Donor Advised Funds': ['donor_advised_funds_cnt', 'donor_advised_funds_value_eoy',
                                    'donor_advised_funds_contri_amt', 'donor_advised_funds_grants_amt'],
        },
        'descriptions': {
            'id': 'Auto-generated primary key',
            'np_return_id': 'FK to nonprofit_returns.id',
            'ein': 'Employer Identification Number',
            'tax_year': 'Tax year of the filing',
            'endowment_boy_amt': 'Endowment balance at beginning of current year',
            'contributions_amt': 'Contributions to endowment during current year',
            'investment_earnings_amt': 'Net investment earnings/losses during current year',
            'grants_or_scholarships_amt': 'Grants or scholarships paid from endowment',
            'other_expenditures_amt': 'Other expenditures from endowment',
            'admin_expenses_amt': 'Administrative expenses charged to endowment',
            'endowment_eoy_amt': 'Endowment balance at end of current year',
            'board_designated_pct': 'Percentage that is board-designated (quasi-endowment)',
            'perm_restricted_pct': 'Percentage that is permanently restricted',
            'temp_restricted_pct': 'Percentage that is temporarily restricted',
            'yr1_endowment_boy_amt': 'Year -1: endowment beginning balance',
            'yr1_contributions_amt': 'Year -1: contributions',
            'yr1_investment_earnings_amt': 'Year -1: investment earnings/losses',
            'yr1_grants_or_scholarships_amt': 'Year -1: grants or scholarships',
            'yr1_other_expenditures_amt': 'Year -1: other expenditures',
            'yr1_admin_expenses_amt': 'Year -1: admin expenses',
            'yr1_endowment_eoy_amt': 'Year -1: endowment ending balance',
            'yr2_endowment_boy_amt': 'Year -2: endowment beginning balance',
            'yr2_contributions_amt': 'Year -2: contributions',
            'yr2_investment_earnings_amt': 'Year -2: investment earnings/losses',
            'yr2_grants_or_scholarships_amt': 'Year -2: grants or scholarships',
            'yr2_other_expenditures_amt': 'Year -2: other expenditures',
            'yr2_admin_expenses_amt': 'Year -2: admin expenses',
            'yr2_endowment_eoy_amt': 'Year -2: endowment ending balance',
            'yr3_endowment_boy_amt': 'Year -3: endowment beginning balance',
            'yr3_contributions_amt': 'Year -3: contributions',
            'yr3_investment_earnings_amt': 'Year -3: investment earnings/losses',
            'yr3_grants_or_scholarships_amt': 'Year -3: grants or scholarships',
            'yr3_other_expenditures_amt': 'Year -3: other expenditures',
            'yr3_admin_expenses_amt': 'Year -3: admin expenses',
            'yr3_endowment_eoy_amt': 'Year -3: endowment ending balance',
            'yr4_endowment_boy_amt': 'Year -4: endowment beginning balance',
            'yr4_contributions_amt': 'Year -4: contributions',
            'yr4_investment_earnings_amt': 'Year -4: investment earnings/losses',
            'yr4_grants_or_scholarships_amt': 'Year -4: grants or scholarships',
            'yr4_other_expenditures_amt': 'Year -4: other expenditures',
            'yr4_admin_expenses_amt': 'Year -4: admin expenses',
            'yr4_endowment_eoy_amt': 'Year -4: endowment ending balance',
            'donor_advised_funds_cnt': 'Number of donor advised funds',
            'donor_advised_funds_value_eoy': 'Aggregate value of donor advised funds at EOY',
            'donor_advised_funds_contri_amt': 'Total contributions to donor advised funds',
            'donor_advised_funds_grants_amt': 'Total grants from donor advised funds',
            'created_at': 'When this row was inserted into our DB',
        },
    },
}


def get_connection():
    return psycopg2.connect(**DB_CONFIG)


def get_table_stats(conn, table_name):
    """Get row count and unique EIN count for a table."""
    cur = conn.cursor()
    cur.execute(f"SELECT count(*) FROM f990_2025.{table_name}")
    total_rows = cur.fetchone()[0]

    # Find the EIN column name
    ein_col = None
    for candidate in ['ein', 'filer_ein']:
        cur.execute("""
            SELECT column_name FROM information_schema.columns
            WHERE table_schema = 'f990_2025' AND table_name = %s AND column_name = %s
        """, (table_name, candidate))
        if cur.fetchone():
            ein_col = candidate
            break

    unique_eins = None
    if ein_col:
        cur.execute(f"SELECT count(DISTINCT {ein_col}) FROM f990_2025.{table_name}")
        unique_eins = cur.fetchone()[0]

    cur.close()
    return total_rows, unique_eins, ein_col


def get_column_populations(conn, table_name, columns):
    """Get population count for each column using a single query."""
    count_exprs = ', '.join(f'count("{c}") AS "{c}"' for c in columns)
    cur = conn.cursor()
    cur.execute(f"SELECT {count_exprs} FROM f990_2025.{table_name}")
    row = cur.fetchone()
    cur.close()
    return dict(zip(columns, row))


def get_column_types(conn, table_name):
    """Get column types from information_schema."""
    cur = conn.cursor()
    cur.execute("""
        SELECT column_name, data_type, udt_name
        FROM information_schema.columns
        WHERE table_schema = 'f990_2025' AND table_name = %s
        ORDER BY ordinal_position
    """, (table_name,))
    results = {}
    for col_name, data_type, udt_name in cur.fetchall():
        type_map = {
            'int4': 'INTEGER',
            'int8': 'BIGINT',
            'varchar': 'VARCHAR',
            'text': 'TEXT',
            'numeric': 'NUMERIC',
            'bool': 'BOOLEAN',
            'date': 'DATE',
            'timestamp': 'TIMESTAMP',
            'jsonb': 'JSONB',
            'float8': 'DOUBLE',
        }
        results[col_name] = type_map.get(udt_name, data_type.upper())
    cur.close()
    return results


def type_label(udt):
    """Friendly type label."""
    return {
        'INTEGER': 'INTEGER', 'BIGINT': 'BIGINT', 'VARCHAR': 'VARCHAR',
        'TEXT': 'TEXT', 'NUMERIC': 'NUMERIC', 'BOOLEAN': 'BOOLEAN',
        'DATE': 'DATE', 'TIMESTAMP': 'TIMESTAMP', 'JSONB': 'JSONB',
        'DOUBLE': 'DOUBLE',
    }.get(udt, udt)


def generate_md(table_name, tdef, total_rows, unique_eins, ein_col, col_types, populations):
    """Generate markdown column reference."""
    lines = []
    lines.append(f"# {table_name} -- Column Reference\n")
    lines.append(f"**Table:** `f990_2025.{table_name}`")
    lines.append(f"**Granularity:** {tdef['granularity']}")
    lines.append(f"**Total rows:** {total_rows:,}")
    if unique_eins is not None:
        lines.append(f"**Unique EINs:** {unique_eins:,}")
    lines.append(f"**Source:** {tdef['source']}")
    lines.append("")
    lines.append("---")
    lines.append("")

    dead_cols = []

    for section_name, col_list in tdef['sections'].items():
        lines.append(f"## {section_name}\n")
        lines.append("| Column | Type | Populated | % | Meaning |")
        lines.append("|--------|------|---:|---:|---------|")

        for col in col_list:
            ctype = col_types.get(col, '?')
            pop = populations.get(col, 0)
            pct = f"{pop / total_rows * 100:.1f}%" if total_rows > 0 and pop > 0 else "0%"
            desc = tdef['descriptions'].get(col, '')
            if pop == 0 and col not in ('created_at',):
                dead_cols.append(col)
            lines.append(f"| `{col}` | {ctype} | {pop:,} | {pct} | {desc} |")

        lines.append("")

    if dead_cols:
        lines.append("---\n")
        lines.append("## Dead Columns (never populated)\n")
        lines.append(f"These columns exist but have zero data across all {total_rows:,} rows:\n")
        for c in dead_cols:
            desc = tdef['descriptions'].get(c, '')
            lines.append(f"- `{c}` -- {desc}")
        lines.append("")

    return '\n'.join(lines)


def generate_docx(table_name, tdef, total_rows, unique_eins, ein_col, col_types, populations):
    """Generate Word document column reference."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(f'{table_name} - Column Reference', level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    # Subtitle info
    p = doc.add_paragraph()
    p.add_run('Table: ').bold = True
    p.add_run(f'f990_2025.{table_name}\n')
    p.add_run('Granularity: ').bold = True
    p.add_run(f'{tdef["granularity"]}\n')
    p.add_run('Total rows: ').bold = True
    p.add_run(f'{total_rows:,}\n')
    if unique_eins is not None:
        p.add_run('Unique EINs: ').bold = True
        p.add_run(f'{unique_eins:,}\n')
    p.add_run('Source: ').bold = True
    p.add_run(tdef['source'])

    dead_cols = []

    for section_name, col_list in tdef['sections'].items():
        doc.add_heading(section_name, level=2)

        table = doc.add_table(rows=1 + len(col_list), cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        headers = ['Column', 'Type', 'Populated', '%', 'Meaning']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): HEADER_BG,
                qn('w:val'): 'clear',
            })
            shading.append(shading_elm)

        # Data rows
        for row_idx, col in enumerate(col_list):
            ctype = col_types.get(col, '?')
            pop = populations.get(col, 0)
            pct = f"{pop / total_rows * 100:.1f}%" if total_rows > 0 and pop > 0 else "0%"
            desc = tdef['descriptions'].get(col, '')
            if pop == 0 and col not in ('created_at',):
                dead_cols.append(col)

            values = [col, ctype, f"{pop:,}", pct, desc]
            for col_idx, val in enumerate(values):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
                if pop == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.2)
            row.cells[1].width = Inches(0.8)
            row.cells[2].width = Inches(0.9)
            row.cells[3].width = Inches(0.5)
            row.cells[4].width = Inches(3.1)

        doc.add_paragraph()  # spacer

    if dead_cols:
        doc.add_heading('Dead Columns (never populated)', level=2)
        p = doc.add_paragraph()
        p.add_run(f'These columns exist but have zero data across all {total_rows:,} rows: ')
        p.add_run(', '.join(dead_cols))

    return doc


def main():
    conn = get_connection()
    print(f"Connected to database. Processing {len(TABLE_DEFS)} tables...\n")

    for table_name, tdef in TABLE_DEFS.items():
        print(f"  {table_name}...")

        # Get stats
        total_rows, unique_eins, ein_col = get_table_stats(conn, table_name)
        col_types = get_column_types(conn, table_name)
        all_cols = list(col_types.keys())
        populations = get_column_populations(conn, table_name, all_cols)

        # Generate markdown
        md_content = generate_md(table_name, tdef, total_rows, unique_eins, ein_col, col_types, populations)
        md_path = os.path.join(OUTPUT_DIR, f'{table_name}_column_reference.md')
        with open(md_path, 'w') as f:
            f.write(md_content)

        # Generate docx
        doc = generate_docx(table_name, tdef, total_rows, unique_eins, ein_col, col_types, populations)
        docx_path = os.path.join(OUTPUT_DIR, f'{table_name}_column_reference.docx')
        doc.save(docx_path)

        print(f"    -> {total_rows:,} rows, {len(all_cols)} columns -> .md + .docx")

    conn.close()
    print(f"\nDone! Generated {len(TABLE_DEFS) * 2} files in {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
