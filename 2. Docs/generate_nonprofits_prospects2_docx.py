#!/usr/bin/env python3
"""Generate Word doc for nonprofits_prospects2 column reference."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

TOTAL = 673381
HEADER_BG = '1e3a5f'


def styled_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn('w:shd'), {
            qn('w:fill'): HEADER_BG, qn('w:val'): 'clear'}))


def add_column_table(doc, section_name, rows, headers=None):
    if headers is None:
        headers = ['Column', 'Type', 'Populated', '%', 'Meaning']
    doc.add_heading(section_name, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    styled_header(table, headers)
    for row_idx, (col_name, col_type, count, meaning) in enumerate(rows):
        pct = f"{count / TOTAL * 100:.1f}%" if count > 0 else "0%"
        values = [col_name, col_type, f"{count:,}", pct, meaning]
        for col_idx, val in enumerate(values):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if count < TOTAL * 0.15:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(0.4)
        row.cells[4].width = Inches(3.3)
    doc.add_paragraph()


def add_code_table(doc, title, rows, col_headers=None):
    if col_headers is None:
        col_headers = ['Code', 'Meaning']
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    table = doc.add_table(rows=1 + len(rows), cols=len(col_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    styled_header(table, col_headers)
    for row_idx, vals in enumerate(rows):
        for col_idx, val in enumerate(vals):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


# ── Title ──────────────────────────────────────────────────────────────

title = doc.add_heading('nonprofits_prospects2 - Column Reference', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

p = doc.add_paragraph()
p.add_run('Table: ').bold = True
p.add_run('f990_2025.nonprofits_prospects2\n')
p.add_run('Granularity: ').bold = True
p.add_run('One row per nonprofit organization (EIN is primary key)\n')
p.add_run('Total rows: ').bold = True
p.add_run('673,381\n')
p.add_run('Sources: ').bold = True
p.add_run('nonprofit_returns (990/990-EZ XML filings) + bmf (IRS Business Master File)\n')
p.add_run('Created: ').bold = True
p.add_run('2026-02-11')

# ── Column tables ──────────────────────────────────────────────────────

add_column_table(doc, "Identity & Filing (from nonprofit_returns)", [
    ("id", "INTEGER", 673381, "Sequential ID (ROW_NUMBER over EIN)"),
    ("ein", "VARCHAR", 673381, "Employer Identification Number (9-digit, no dashes)"),
    ("organization_name", "TEXT", 673381, "Nonprofit's legal name from most recent return"),
    ("most_recent_tax_year", "INTEGER", 673381, "Most recent tax year with a filing"),
    ("most_recent_tax_period_begin", "DATE", 673381, "Start of most recent fiscal year"),
    ("most_recent_tax_period_end", "DATE", 673381, "End of most recent fiscal year"),
    ("most_recent_form_type", "VARCHAR", 673381, "IRS form: 990 (larger orgs) or 990EZ (smaller)"),
])

add_column_table(doc, "Address & Contact (from nonprofit_returns)", [
    ("address_line1", "TEXT", 672786, "Mailing address. Most recent non-null."),
    ("address_line2", "TEXT", 0, "Suite/floor (dead column)"),
    ("city", "TEXT", 672786, "City (from same row as address_line1)"),
    ("state", "VARCHAR", 672786, "State code (from same row as address_line1)"),
    ("zip", "VARCHAR", 672786, "ZIP code (from same row as address_line1)"),
    ("phone", "VARCHAR", 673381, "Phone number. Most recent non-null."),
    ("website", "TEXT", 578028, "Website URL. Most recent non-null."),
])

add_column_table(doc, "Mission & Programs (from nonprofit_returns)", [
    ("mission_description", "TEXT", 673381, "Mission statement. Most recent non-null. Both 990 and 990-EZ."),
    ("org_description", "TEXT", 673381, "Activity description (990) or primary exempt purpose (990-EZ). COALESCE."),
    ("program_1_desc", "TEXT", 377992, "Program #1 description (Part III Line 4a). Primarily 990 filers."),
    ("program_2_desc", "TEXT", 161782, "Program #2 description (Part III Line 4b)"),
    ("program_3_desc", "TEXT", 81064, "Program #3 description (Part III Line 4c)"),
    ("program_1_expense_amt", "NUMERIC", 295044, "Expenses for program #1"),
    ("program_2_expense_amt", "NUMERIC", 138564, "Expenses for program #2"),
    ("program_3_expense_amt", "NUMERIC", 70443, "Expenses for program #3"),
])

add_column_table(doc, "Classification & Staffing (from nonprofit_returns)", [
    ("ntee_code", "VARCHAR", 452209, "NTEE code from 990 filing (see NTEE table below)"),
    ("total_employees_cnt", "INTEGER", 389134, "Total employees (primarily 990 filers)"),
    ("total_volunteers_cnt", "INTEGER", 322949, "Total volunteers"),
])

add_column_table(doc, "BMF Enrichment (from IRS Business Master File)", [
    ("bmf_status", "VARCHAR(5)", 604428, "Exempt status: 01=active, 12=revoked, 25=terminated, NULL=not in BMF"),
    ("bmf_ico", "TEXT", 369722, "In Care Of name (contact person at org)"),
    ("bmf_subsection", "VARCHAR(5)", 604428, "IRC subsection (see code tables below)"),
    ("bmf_foundation", "VARCHAR(5)", 604428, "Foundation status code (see code tables below)"),
    ("bmf_deductibility", "VARCHAR(5)", 604428, "Contribution deductibility (see code tables below)"),
    ("bmf_ruling", "VARCHAR(6)", 604428, "IRS ruling date (YYYYMM format)"),
    ("bmf_ntee_cd", "VARCHAR(10)", 450409, "NTEE code from BMF (compare with ntee_code from 990 filing)"),
    ("bmf_org_type", "VARCHAR(5)", 604428, "Legal form (see code tables below)"),
    ("bmf_affiliation", "VARCHAR(5)", 604428, "Organizational relationship (see code tables below)"),
    ("bmf_group_code", "VARCHAR(10)", 604428, "Group exemption number. 0000 = not part of a group."),
    ("bmf_classification", "VARCHAR(10)", 604428, "Legacy IRS classification codes"),
    ("bmf_activity", "VARCHAR(10)", 604428, "Legacy activity codes. Superseded by NTEE."),
])

# ── Code tables ────────────────────────────────────────────────────────

doc.add_heading('BMF Code Tables', level=1)
doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

add_code_table(doc, 'Subsection Codes (bmf_subsection)', [
    ('03', '501(c)(3) - Charitable, educational, religious, scientific. Tax-deductible donations.'),
    ('04', '501(c)(4) - Civic leagues, social welfare. Example: Sierra Club, NRA.'),
    ('05', '501(c)(5) - Labor, agricultural, horticultural. Example: AFL-CIO locals.'),
    ('06', '501(c)(6) - Business leagues, chambers of commerce.'),
    ('07', '501(c)(7) - Social and recreational clubs. Example: country clubs.'),
    ('08', '501(c)(8) - Fraternal beneficiary societies. Example: Elks lodges.'),
    ('09', '501(c)(9) - Voluntary employees beneficiary associations (VEBAs).'),
    ('10', '501(c)(10) - Domestic fraternal societies.'),
    ('12', '501(c)(12) - Benevolent life insurance, mutual telephone/electric companies.'),
    ('13', '501(c)(13) - Cemetery companies.'),
    ('14', '501(c)(14) - Credit unions, mutual reserve funds.'),
    ('19', '501(c)(19) - Veterans organizations. Example: VFW, American Legion.'),
    ('25', '501(c)(25) - Title holding corps for multiple exempt orgs.'),
    ('92', '527 - Political organizations, PACs.'),
])

add_code_table(doc, 'Foundation Codes (bmf_foundation)', [
    ('00', 'Not a 501(c)(3) organization (field not applicable).'),
    ('02', 'Private operating foundation exempt from excise taxes on net investment income.'),
    ('03', 'Private operating foundation (all other).'),
    ('04', 'Private non-operating foundation (the classic grantmaking PF).'),
    ('10', 'Church (170(b)(1)(A)(i)).'),
    ('11', 'School (170(b)(1)(A)(ii)).'),
    ('12', 'Hospital or cooperative hospital service org (170(b)(1)(A)(iii)).'),
    ('13', 'Org operated for benefit of a college or university (170(b)(1)(A)(iv)).'),
    ('14', 'Governmental unit (170(b)(1)(A)(v)).'),
    ('15', 'Publicly supported, >1/3 from contributions (509(a)(1)). Example: United Way.'),
    ('16', 'Publicly supported, >1/3 from gross receipts (509(a)(2)). Example: YMCA.'),
    ('17', 'Organization determined publicly supported by IRS (alternative test).'),
    ('21', 'Supporting organization Type I (509(a)(3)). Operated by supported org.'),
    ('22', 'Supporting organization Type II (509(a)(3)). Controlled in connection with.'),
    ('23', 'Supporting organization Type III functionally integrated (509(a)(3)).'),
    ('24', 'Supporting organization Type III other (509(a)(3)).'),
])

add_code_table(doc, 'Status Codes (bmf_status)', [
    ('01', 'Unconditional Exemption. Active and in good standing.'),
    ('02', 'Conditional Exemption. Exempt with conditions.'),
    ('12', 'Revoked. Usually auto-revocation for 3 years non-filing.'),
    ('25', 'Terminated. Voluntarily or IRS-terminated.'),
    ('NULL', 'Not found in current BMF. Likely dissolved, merged, or foreign.'),
])

add_code_table(doc, 'Deductibility Codes (bmf_deductibility)', [
    ('1', 'Contributions are deductible. Standard charitable deduction.'),
    ('2', 'Contributions deductible by treaty or election. Limited.'),
    ('0', 'Contributions NOT deductible. Example: social clubs, business leagues.'),
    ('4', 'Deductibility depends on date of gift.'),
])

add_code_table(doc, 'Organization Type Codes (bmf_org_type)', [
    ('1', 'Corporation. Most common form for nonprofits.'),
    ('2', 'Trust. Common for private foundations.'),
    ('3', 'Co-operative. Member-owned organizations.'),
    ('4', 'Partnership. Rare for exempt orgs.'),
    ('5', 'Association. Unincorporated membership orgs. Example: churches, VFW posts.'),
    ('6', 'Other.'),
    ('0', 'Not specified or unknown.'),
])

add_code_table(doc, 'Affiliation Codes (bmf_affiliation)', [
    ('3', 'Independent. Not part of a group exemption. The majority.'),
    ('9', 'Subordinate. Part of a group exemption. Example: local Girl Scout council.'),
    ('1', 'Central organization (group ruling). Parent holding group exemption.'),
    ('6', 'Central org NOT included in its own group return.'),
    ('0', 'Not specified.'),
    ('2', 'Intermediate org (group return). Files group return with subordinates.'),
    ('8', 'Central org whose group exemption has been revoked.'),
    ('7', 'Intermediate org (not filing group return).'),
])

add_code_table(doc, 'NTEE Major Categories (first letter of bmf_ntee_cd or ntee_code)', [
    ('A', 'Arts, Culture, Humanities. Example: A20=Museums, A60=Performing Arts.'),
    ('B', 'Education. Example: B20=Elementary/Secondary, B40=Higher Ed.'),
    ('C', 'Environment. Example: C20=Pollution Abatement, C30=Natural Resources.'),
    ('D', 'Animal-Related. Example: D20=Animal Protection, D30=Wildlife.'),
    ('E', 'Health Care. Example: E20=Hospitals, E40=Reproductive Health.'),
    ('F', 'Mental Health & Crisis. Example: F20=Substance Abuse.'),
    ('G', 'Diseases & Disorders. Example: G20=Birth Defects, G40=Cancer.'),
    ('H', 'Medical Research. Example: H20=Birth Defects Research.'),
    ('I', 'Crime & Legal. Example: I20=Crime Prevention.'),
    ('J', 'Employment. Example: J20=Job Training.'),
    ('K', 'Food, Agriculture, Nutrition. Example: K20=Food Banks.'),
    ('L', 'Housing & Shelter. Example: L20=Housing Development.'),
    ('M', 'Public Safety & Disaster. Example: M20=Disaster Preparedness.'),
    ('N', 'Recreation & Sports. Example: N20=Camps, N60=Amateur Sports.'),
    ('O', 'Youth Development. Example: O20=Youth Centers, O40=Scouting.'),
    ('P', 'Human Services. Example: P20=Multipurpose, P80=Emergency Assistance.'),
    ('Q', 'International Affairs.'),
    ('R', 'Civil Rights & Social Action.'),
    ('S', 'Community Improvement. Example: S20=Coalitions, S40=Business.'),
    ('T', 'Philanthropy & Grantmaking. Example: T20=Private Foundations.'),
    ('U', 'Science & Technology.'),
    ('V', 'Social Science.'),
    ('W', 'Public/Society Benefit.'),
    ('X', 'Religion-Related. Example: X20=Christian, X30=Jewish.'),
    ('Y', 'Mutual & Membership Benefit.'),
    ('Z', 'Unknown / Unclassified.'),
])

add_code_table(doc, 'Asset/Income Size Codes (for reference when querying BMF)', [
    ('0', '$0 or not reported'),
    ('1', '$1 to $10,000'),
    ('2', '$10,000 to $25,000'),
    ('3', '$25,000 to $100,000'),
    ('4', '$100,000 to $500,000'),
    ('5', '$500,000 to $1,000,000'),
    ('6', '$1,000,000 to $5,000,000'),
    ('7', '$5,000,000 to $10,000,000'),
    ('8', '$10,000,000 to $50,000,000'),
    ('9', '$50,000,000+'),
])

add_code_table(doc, 'Filing Requirement Codes (for reference when querying BMF)', [
    ('01', 'Form 990 required. Gross receipts >= $200K or assets >= $500K.'),
    ('02', 'Form 990 or 990-EZ. Gross receipts < $200K and assets < $500K.'),
    ('03', 'Form 990-PF required (private foundations).'),
    ('06', 'Not required to file. Churches, government agencies.'),
    ('00', 'Form 990-N (e-Postcard) eligible. Gross receipts <= $50K.'),
    ('07', 'Section 4947(a)(1) nonexempt charitable trust treated as PF.'),
    ('13', 'Section 4947(a)(1) nonexempt charitable trust not treated as PF.'),
    ('14', 'Not required to file (other exemptions).'),
])

output_path = '/Users/aleckleinman/Documents/TheGrantScout/2. Docs/nonprofits_prospects2_column_reference.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
