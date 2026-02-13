#!/usr/bin/env python3
"""Generate Word doc for nonprofit_returns column reference."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

TOTAL = 673381

sections_data = [
    ("Record Metadata", [
        ("id", "INTEGER", 673381, "Auto-generated primary key"),
        ("ein", "VARCHAR", 673381, "Employer Identification Number (9-digit, no dashes)"),
        ("organization_name", "TEXT", 673381, "Nonprofit's legal name as filed"),
        ("tax_year", "INTEGER", 673381, "Tax year this return covers (e.g., 2023)"),
        ("tax_period_begin", "DATE", 673381, "Start date of the fiscal year"),
        ("tax_period_end", "DATE", 673381, "End date of the fiscal year"),
        ("form_type", "VARCHAR", 673381, "IRS form type: 990 (larger orgs) or 990EZ (smaller orgs)"),
        ("source_file", "TEXT", 673381, "Path to the IRS XML file this was parsed from"),
        ("object_id", "VARCHAR", 560119, "IRS internal identifier for this specific filing"),
        ("created_at", "TIMESTAMP", 673381, "When this row was inserted into our DB"),
        ("updated_at", "TIMESTAMP", 673381, "When this row was last modified"),
    ]),
    ("Contact Info (990 Header)", [
        ("address_line1", "TEXT", 672786, "Nonprofit's mailing address, line 1"),
        ("address_line2", "TEXT", 0, "Suite/floor/PO Box, line 2 (never populated)"),
        ("city", "TEXT", 672786, "City"),
        ("state", "VARCHAR", 672786, "State code (2-letter)"),
        ("zip", "VARCHAR", 672786, "ZIP code"),
        ("country", "TEXT", 673381, "Country (usually US)"),
        ("phone", "VARCHAR", 673381, "Main phone number"),
        ("website", "TEXT", 578028, "Organization website URL"),
    ]),
    ("Financials (Part I / Balance Sheet)", [
        ("total_revenue", "NUMERIC", 668551, "Total revenue from all sources"),
        ("total_expenses", "NUMERIC", 665673, "Total expenses"),
        ("total_assets_boy", "NUMERIC", 402488, "Total assets at beginning of year (990 only)"),
        ("total_assets_eoy", "NUMERIC", 412797, "Total assets at end of year (990 only)"),
        ("net_assets_boy", "NUMERIC", 646099, "Net assets at beginning of year"),
        ("net_assets_eoy", "NUMERIC", 665951, "Net assets at end of year"),
        ("contributions_grants", "NUMERIC", 639130, "Contributions, gifts, and grants received"),
        ("program_service_revenue", "NUMERIC", 557793, "Revenue from program services (fees, contracts)"),
        ("investment_income", "NUMERIC", 564174, "Investment income (dividends, interest, gains)"),
    ]),
    ("Mission / Purpose / Programs (Part III)", [
        ("mission_description", "TEXT", 673381, "Brief mission statement (both 990 and 990-EZ)"),
        ("activity_description", "TEXT", 412797, "Most significant activities (990 only, Part III Line 1)"),
        ("primary_exempt_purpose", "TEXT", 334254, "Primary exempt purpose (990-EZ only, Part III)"),
        ("program_1_desc", "TEXT", 377992, "Program accomplishment #1 (Part III Line 4a, primarily 990)"),
        ("program_2_desc", "TEXT", 161782, "Program accomplishment #2 (Part III Line 4b)"),
        ("program_3_desc", "TEXT", 81064, "Program accomplishment #3 (Part III Line 4c)"),
        ("program_1_expense_amt", "NUMERIC", 295044, "Expenses for program #1"),
        ("program_2_expense_amt", "NUMERIC", 138564, "Expenses for program #2"),
        ("program_3_expense_amt", "NUMERIC", 70443, "Expenses for program #3"),
        ("program_1_revenue_amt", "NUMERIC", 185875, "Revenue for program #1"),
    ]),
    ("Classification Codes", [
        ("activity_code_1", "VARCHAR", 919, "IRS activity code #1 (rarely populated)"),
        ("activity_code_2", "VARCHAR", 281, "IRS activity code #2 (rarely populated)"),
        ("activity_code_3", "VARCHAR", 173, "IRS activity code #3 (rarely populated)"),
        ("ntee_code", "VARCHAR", 452209, "NTEE code (e.g., B20 = elementary/secondary ed)"),
        ("ruling_date", "DATE", 0, "Date IRS granted tax-exempt status (never populated)"),
    ]),
    ("Governance (Part VI)", [
        ("foreign_activities_ind", "BOOLEAN", 673381, "Activities outside the US?"),
        ("voting_members_governing_body_cnt", "INTEGER", 412797, "Voting members on governing board (990 only)"),
        ("voting_members_independent_cnt", "INTEGER", 412797, "Independent voting members (990 only)"),
        ("total_employees_cnt", "INTEGER", 412797, "Total employees (990 only)"),
        ("total_volunteers_cnt", "INTEGER", 322949, "Total volunteers"),
    ]),
]

# Title
title = doc.add_heading('nonprofit_returns - Column Reference', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

# Subtitle info
p = doc.add_paragraph()
p.add_run('Table: ').bold = True
p.add_run('f990_2025.nonprofit_returns\n')
p.add_run('Granularity: ').bold = True
p.add_run('One row per 990/990-EZ filing (multiple rows per nonprofit)\n')
p.add_run('Total rows: ').bold = True
p.add_run('2,953,274\n')
p.add_run('Unique nonprofits: ').bold = True
p.add_run('673,381\n')
p.add_run('Tax years: ').bold = True
p.add_run('2016-2024\n')
p.add_run('Source: ').bold = True
p.add_run('IRS Form 990 and Form 990-EZ XML e-filings')

HEADER_COLOR = RGBColor(0x1e, 0x3a, 0x5f)
HEADER_BG = '1e3a5f'

for section_name, rows in sections_data:
    doc.add_heading(section_name, level=2)

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    headers = ['Column', 'Type', 'Nonprofits', '%', 'Meaning']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): HEADER_BG,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for row_idx, (col_name, col_type, count, meaning) in enumerate(rows):
        pct = f"{count / TOTAL * 100:.1f}%" if count > 0 else "0%"
        values = [col_name, col_type, f"{count:,}", pct, meaning]
        for col_idx, val in enumerate(values):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            # Gray out dead columns
            if count == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.9)
        row.cells[3].width = Inches(0.5)
        row.cells[4].width = Inches(3.1)

    doc.add_paragraph()  # spacer

# Dead columns note
doc.add_heading('Dead Columns (never populated)', level=2)
p = doc.add_paragraph()
p.add_run('These columns exist but have zero data across all 2.95M rows: ')
dead = ['address_line2', 'ruling_date']
p.add_run(', '.join(dead))

doc.add_paragraph()

# Nearly dead columns note
doc.add_heading('Nearly Dead Columns (<1% populated)', level=2)
p = doc.add_paragraph()
p.add_run('activity_code_1/2/3').bold = True
p.add_run(' -- IRS activity codes. Only 919 / 281 / 173 EINs respectively. These legacy codes are rarely filed in modern e-file returns.')

doc.add_paragraph()

# Form comparison note
doc.add_heading('Form 990 vs 990-EZ Field Availability', level=2)

comp_table = doc.add_table(rows=11, cols=3)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.LEFT

comp_headers = ['Field Category', 'Form 990', 'Form 990-EZ']
for i, h in enumerate(comp_headers):
    cell = comp_table.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): HEADER_BG,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

comp_data = [
    ("Basic info (name, EIN, address, phone)", "Yes", "Yes"),
    ("Mission description", "Yes", "Yes"),
    ("Activity description", "Yes", "No"),
    ("Primary exempt purpose", "No", "Yes"),
    ("Program descriptions (1-3)", "Yes", "Limited"),
    ("Total revenue / expenses", "Yes", "Yes"),
    ("Total assets (BOY/EOY)", "Yes", "No"),
    ("Net assets (BOY/EOY)", "Yes", "Yes"),
    ("Employee / volunteer counts", "Yes", "Volunteers only (partial)"),
    ("Voting board members", "Yes", "No"),
]

for row_idx, (field, f990, f990ez) in enumerate(comp_data):
    for col_idx, val in enumerate([field, f990, f990ez]):
        cell = comp_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

for row in comp_table.rows:
    row.cells[0].width = Inches(3.0)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(2.5)

output_path = '/Users/aleckleinman/Documents/TheGrantScout/2. Docs/nonprofit_returns_column_reference.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
