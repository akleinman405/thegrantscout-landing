#!/usr/bin/env python3
"""Generate Word doc for pf_returns column reference."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

TOTAL = 143184

sections_data = [
    ("Record Metadata", [
        ("id", "INTEGER", 143184, "Auto-generated primary key"),
        ("ein", "VARCHAR", 143184, "Employer Identification Number (9-digit, no dashes)"),
        ("business_name", "TEXT", 143184, "Foundation's legal name as filed"),
        ("tax_year", "INTEGER", 143184, "Tax year this return covers"),
        ("tax_period_begin", "DATE", 143184, "Start date of the fiscal year"),
        ("tax_period_end", "DATE", 143184, "End date of the fiscal year"),
        ("source_file", "TEXT", 143184, "Path to the IRS XML file"),
        ("return_timestamp", "TIMESTAMP", 143184, "When the return was filed to IRS"),
        ("object_id", "VARCHAR", 143184, "IRS internal identifier for this filing"),
        ("created_at", "TIMESTAMP", 143184, "When row was inserted into our DB"),
        ("updated_at", "TIMESTAMP", 143184, "When row was last modified"),
    ]),
    ("Contact Info (990-PF Header)", [
        ("address_line1", "TEXT", 143184, "Foundation's mailing address, line 1"),
        ("address_line2", "TEXT", 0, "Suite/floor/PO Box, line 2 (never populated)"),
        ("city", "TEXT", 143184, "City"),
        ("state", "VARCHAR", 143145, "State code (2-letter)"),
        ("zip", "VARCHAR", 143172, "ZIP code"),
        ("country", "TEXT", 143184, "Country (usually US)"),
        ("phone_num", "VARCHAR", 138009, "Main phone. Enriched: 96% XML, 4% propagated"),
        ("email_address_txt", "TEXT", 0, "Email address (never populated)"),
        ("website_url", "TEXT", 130796, "Website. Enriched: 29% XML, rest scraped"),
    ]),
    ("Foundation Classification (Part VII-A)", [
        ("private_operating_foundation_ind", "BOOLEAN", 143184, "Is a private OPERATING foundation?"),
        ("exempt_operating_foundations_ind", "BOOLEAN", 143184, "Has IRS exempt operating ruling?"),
        ("grants_to_organizations_ind", "BOOLEAN", 143184, "Makes grants to organizations?"),
        ("grants_to_individuals_ind", "BOOLEAN", 143184, "Makes grants to individuals?"),
        ("only_contri_to_preselected_ind", "BOOLEAN", 143184, "Only preselected recipients? FALSE = accepts apps"),
        ("foreign_activities_ind", "BOOLEAN", 143184, "Activities outside the US?"),
        ("contributing_manager_nm", "TEXT", 57118, "Foundation manager name (Part VIII)"),
        ("formation_yr", "INTEGER", 0, "Year formed (never populated)"),
    ]),
    ("Financials (Part I / Balance Sheet)", [
        ("total_revenue_amt", "NUMERIC", 143184, "Total revenue for the year"),
        ("total_expenses_amt", "NUMERIC", 143184, "Total expenses for the year"),
        ("total_assets_eoy_amt", "NUMERIC", 143184, "Total assets at end of year (main size metric)"),
        ("total_assets_boy_amt", "NUMERIC", 138733, "Total assets at beginning of year"),
        ("fmv_assets_eoy_amt", "NUMERIC", 143184, "Fair market value of assets at EOY"),
        ("net_assets_eoy_amt", "NUMERIC", 143184, "Net assets at end of year"),
        ("net_assets_boy_amt", "NUMERIC", 137119, "Net assets at beginning of year"),
        ("total_liabilities_eoy_amt", "NUMERIC", 143184, "Total liabilities at end of year"),
        ("excess_revenue_over_expenses_amt", "NUMERIC", 143184, "Revenue minus expenses (surplus/deficit)"),
        ("contributions_received_amt", "NUMERIC", 89517, "Contributions and gifts received"),
        ("investment_income_amt", "NUMERIC", 143184, "Investment income (dividends, interest, gains)"),
    ]),
    ("Grantmaking (Part XV / Distribution)", [
        ("total_grant_paid_amt", "NUMERIC", 124782, "Total grants paid during this year"),
        ("qualifying_distributions_amt", "NUMERIC", 133860, "Qualifying distributions (toward 5% payout)"),
        ("distributable_as_adjusted_amt", "NUMERIC", 133636, "Required distribution amount (5% rule)"),
        ("min_investment_return_amt", "NUMERIC", 140683, "Minimum investment return (5% of avg FMV)"),
        ("undistributed_income_cy_amt", "NUMERIC", 134030, "Undistributed income, current year"),
        ("undistributed_income_py_ind", "BOOLEAN", 143184, "Has undistributed income from prior years?"),
        ("total_grant_approved_future_amt", "NUMERIC", 66552, "Grants approved but not yet paid"),
        ("grants_payable_eoy_amt", "NUMERIC", 13373, "Outstanding grant obligations at EOY"),
    ]),
    ("Application Process (Part XV - Supplementary)", [
        ("application_submission_info", "JSONB", 35385, "Structured contact: {name, phone, email, address}"),
        ("app_contact_name", "TEXT", 35162, "Contact person for applications"),
        ("app_contact_phone", "VARCHAR", 33619, "Contact phone for applications"),
        ("app_form_requirements", "TEXT", 33301, "How to apply (LETTER, WRITTEN REQUEST, etc.)"),
        ("app_restrictions", "TEXT", 33171, "Restrictions on who can apply"),
        ("app_submission_deadlines", "TEXT", 33066, "When applications are due"),
        ("app_contact_email", "TEXT", 10092, "Contact email for applications"),
    ]),
    ("Mission / Purpose (Part IX-A)", [
        ("activity_or_mission_desc", "TEXT", 115205, "Description of activities or mission"),
        ("mission_desc", "TEXT", 0, "Mission statement (never populated)"),
        ("primary_exempt_purpose", "TEXT", 0, "Primary exempt purpose (never populated)"),
        ("activity_code_1", "VARCHAR", 0, "IRS activity code #1 (never populated)"),
        ("activity_code_2", "VARCHAR", 0, "IRS activity code #2 (never populated)"),
        ("activity_code_3", "VARCHAR", 0, "IRS activity code #3 (never populated)"),
    ]),
    ("Enriched Fields (added post-import)", [
        ("ntee_code", "VARCHAR", 90868, "NTEE classification code (e.g., T20)"),
        ("ntee_source", "VARCHAR", 122430, "Source of NTEE code (always irs_bmf)"),
        ("ntee_enriched_at", "TIMESTAMP", 122430, "When NTEE lookup was done"),
        ("url_source", "VARCHAR", 124814, "How website URL was found"),
        ("url_enriched_at", "TIMESTAMP", 124814, "When URL enrichment was done"),
        ("url_validated", "BOOLEAN", 124814, "Whether URL was checked and working"),
        ("phone_source", "VARCHAR", 117703, "How phone was obtained (xml or propagated)"),
        ("phone_enriched_at", "TIMESTAMP", 117703, "When phone enrichment was done"),
        ("phone_validated", "BOOLEAN", 117703, "Whether phone number was verified"),
    ]),
]

# Title
title = doc.add_heading('pf_returns - Column Reference', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

# Subtitle info
p = doc.add_paragraph()
p.add_run('Table: ').bold = True
p.add_run('f990_2025.pf_returns\n')
p.add_run('Granularity: ').bold = True
p.add_run('One row per 990-PF filing (multiple rows per foundation)\n')
p.add_run('Total rows: ').bold = True
p.add_run('638,698\n')
p.add_run('Unique foundations: ').bold = True
p.add_run('143,184\n')
p.add_run('Tax years: ').bold = True
p.add_run('2016-2024')

HEADER_COLOR = RGBColor(0x1e, 0x3a, 0x5f)
HEADER_BG = '1e3a5f'

for section_name, rows in sections_data:
    doc.add_heading(section_name, level=2)

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    headers = ['Column', 'Type', 'Foundations', '%', 'Meaning']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): HEADER_BG,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for row_idx, (col_name, col_type, count, meaning) in enumerate(rows):
        pct = f"{count / TOTAL * 100:.1f}%" if count > 0 else "0%"
        values = [col_name, col_type, f"{count:,}", pct, meaning]
        for col_idx, val in enumerate(values):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            # Gray out dead columns
            if count == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.9)
        row.cells[3].width = Inches(0.5)
        row.cells[4].width = Inches(3.1)

    doc.add_paragraph()  # spacer

# Dead columns note
doc.add_heading('Dead Columns (never populated)', level=2)
p = doc.add_paragraph()
p.add_run('These columns exist but have zero data across all 638K rows: ')
dead = ['address_line2', 'email_address_txt', 'formation_yr', 'mission_desc',
        'primary_exempt_purpose', 'activity_code_1/2/3']
p.add_run(', '.join(dead))

output_path = '/Users/aleckleinman/Documents/TheGrantScout/2. Docs/pf_returns_column_reference.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
