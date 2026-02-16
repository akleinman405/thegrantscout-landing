#!/usr/bin/env python3
"""
TheGrantScout Markdown → Word Converter
========================================

Converts markdown reports to professionally styled Word documents.

Usage:
    python3 "0. Tools/md_to_docx.py" --input report.md --output report.docx
    python3 "0. Tools/md_to_docx.py" --input report.md  # output = same dir, .docx extension

Brand: Navy (#1e3a5f) + Gold (#d4a853) + Charcoal (#2C3E50)
Fonts: Calibri Light (headings), Calibri (body)
Based on 09_convert_to_docx.py (gold standard, 744 lines).
"""

import argparse
import os
import re
import sys

# Allow importing branding.py from same directory
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import branding

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# BRAND COLORS (from branding.py → python-docx RGBColor)
# ============================================================================

NAVY = RGBColor(*branding.NAVY_RGB)
NAVY_DARK = RGBColor(*branding.NAVY_DARK_RGB)
GOLD = RGBColor(*branding.GOLD_RGB)
GOLD_DARK = RGBColor(*branding.GOLD_DARK_RGB)
CHARCOAL = RGBColor(*branding.CHARCOAL_RGB)
GRAY = RGBColor(*branding.GRAY_RGB)
WHITE = RGBColor(*branding.WHITE_RGB)
SUCCESS = RGBColor(*branding.SUCCESS_RGB)
WARNING = RGBColor(*branding.WARNING_RGB)
ERROR = RGBColor(*branding.ERROR_RGB)

# Background hex codes (bare, for XML)
BG_NAVY = branding.NAVY_BARE
BG_WHITE = branding.WHITE_BARE
BG_LIGHT_GRAY = branding.BG_LIGHT_GRAY_BARE
BG_LIGHT_GOLD = branding.BG_LIGHT_GOLD_BARE
BG_LIGHT_NAVY = branding.BG_LIGHT_NAVY_BARE
BG_BORDER_LIGHT = branding.BG_BORDER_LIGHT_BARE

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_bottom_border(cell, color_hex='d4a853', size='12'):
    """Add bottom border to cell (for gold underline effect on headers)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:color'), color_hex)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def create_styled_document():
    """Create document with TheGrantScout professional styles."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = doc.styles

    # Title (32pt Bold Navy)
    title_style = styles.add_style('TGS Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = branding.FONT_HEADING
    title_style.font.size = Pt(branding.SIZE_TITLE)
    title_style.font.bold = True
    title_style.font.color.rgb = NAVY
    title_style.paragraph_format.space_after = Pt(0)
    title_style.paragraph_format.space_before = Pt(0)

    # Subtitle (18pt Gray)
    subtitle_style = styles.add_style('TGS Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = branding.FONT_HEADING
    subtitle_style.font.size = Pt(branding.SIZE_SUBTITLE)
    subtitle_style.font.color.rgb = GRAY
    subtitle_style.paragraph_format.space_after = Pt(8)
    subtitle_style.paragraph_format.space_before = Pt(0)

    # Metadata (11pt Gray)
    meta_style = styles.add_style('TGS Meta', WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.name = branding.FONT_BODY
    meta_style.font.size = Pt(branding.SIZE_BODY)
    meta_style.font.color.rgb = GRAY
    meta_style.paragraph_format.space_after = Pt(4)

    # Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.name = branding.FONT_HEADING
    h1_style.font.size = Pt(branding.SIZE_H1)
    h1_style.font.bold = True
    h1_style.font.color.rgb = NAVY
    h1_style.paragraph_format.space_before = Pt(branding.SPACE_BEFORE_H1)
    h1_style.paragraph_format.space_after = Pt(branding.SPACE_AFTER_H1)

    # Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.name = branding.FONT_BODY
    h2_style.font.size = Pt(branding.SIZE_H2)
    h2_style.font.bold = True
    h2_style.font.color.rgb = NAVY_DARK
    h2_style.paragraph_format.space_before = Pt(branding.SPACE_BEFORE_H2)
    h2_style.paragraph_format.space_after = Pt(branding.SPACE_AFTER_H2)

    # Heading 3
    h3_style = styles['Heading 3']
    h3_style.font.name = branding.FONT_BODY
    h3_style.font.size = Pt(branding.SIZE_H3)
    h3_style.font.bold = True
    h3_style.font.color.rgb = CHARCOAL
    h3_style.paragraph_format.space_before = Pt(branding.SPACE_BEFORE_H3)
    h3_style.paragraph_format.space_after = Pt(branding.SPACE_AFTER_H3)

    # Custom heading aliases (backwards compat)
    h1_custom = styles.add_style('TGS H1', WD_STYLE_TYPE.PARAGRAPH)
    h1_custom.base_style = h1_style
    h2_custom = styles.add_style('TGS H2', WD_STYLE_TYPE.PARAGRAPH)
    h2_custom.base_style = h2_style
    h3_custom = styles.add_style('TGS H3', WD_STYLE_TYPE.PARAGRAPH)
    h3_custom.base_style = h3_style

    # Body text (11pt Charcoal, 1.15 line spacing)
    body_style = styles.add_style('TGS Body', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = branding.FONT_BODY
    body_style.font.size = Pt(branding.SIZE_BODY)
    body_style.font.color.rgb = CHARCOAL
    body_style.paragraph_format.space_after = Pt(branding.SPACE_AFTER_PARA)
    body_style.paragraph_format.line_spacing = 1.15

    # Urgent status label (11pt Bold Red)
    urgent_style = styles.add_style('TGS Urgent', WD_STYLE_TYPE.PARAGRAPH)
    urgent_style.font.name = branding.FONT_BODY
    urgent_style.font.size = Pt(branding.SIZE_BODY)
    urgent_style.font.bold = True
    urgent_style.font.color.rgb = ERROR
    urgent_style.paragraph_format.space_after = Pt(branding.SPACE_AFTER_PARA)

    return doc


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table and return rows."""
    rows = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                rows.append(cells)
        i += 1
    return rows, i


def add_styled_table(doc, rows):
    """Add a professionally styled table with TGS branding."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True

        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                continue
            cell = row.cells[j]
            clean_text = cell_text.replace('**', '').strip()
            cell.text = clean_text

            if i == 0:
                set_cell_shading(cell, BG_NAVY)
                add_bottom_border(cell, branding.GOLD_BARE, '12')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.keep_with_next = True
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
                        run.font.size = Pt(branding.SIZE_TABLE_HEADER)
                        run.font.name = branding.FONT_BODY
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, BG_LIGHT_GRAY)
                else:
                    set_cell_shading(cell, BG_WHITE)

                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.keep_with_next = True
                    for run in para.runs:
                        run.font.size = Pt(branding.SIZE_TABLE_DATA)
                        run.font.color.rgb = CHARCOAL
                        run.font.name = branding.FONT_BODY

                        upper_text = clean_text.upper()
                        if 'URGENT' in upper_text:
                            run.font.color.rgb = ERROR
                            run.font.bold = True
                        elif 'HIGH' in upper_text:
                            run.font.color.rgb = GOLD_DARK
                            run.font.bold = True

            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.space_before = Pt(4)

    doc.add_paragraph()


def add_callout_box(doc, text):
    """Add a styled callout box with gold left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), BG_LIGHT_GOLD)
    tcPr.append(shading)

    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '32')
    left.set(qn('w:color'), branding.GOLD_BARE)
    tcBorders.append(left)
    for border_name in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), BG_BORDER_LIGHT)
        tcBorders.append(border)
    tcPr.append(tcBorders)

    cell.text = text.replace('**', '')
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.left_indent = Inches(0.1)
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = NAVY
            run.font.name = branding.FONT_BODY

    doc.add_paragraph()


def add_horizontal_rule(doc):
    """Add a gold horizontal rule."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), branding.GOLD_BARE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def clean_markdown_link(text):
    """Convert [text](url) to just text."""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def add_inline_bold(para, text, default_size=None, default_color=None):
    """Add text to paragraph, rendering **bold** segments."""
    if default_size is None:
        default_size = branding.SIZE_BODY
    if default_color is None:
        default_color = CHARCOAL

    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.size = Pt(default_size)
        run.font.color.rgb = default_color
        run.font.name = branding.FONT_BODY


def convert_markdown_to_docx(md_path, docx_path):
    """Convert markdown file to professionally styled Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = create_styled_document()

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    run = header_para.add_run(branding.BRAND_NAME)
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.name = branding.FONT_HEADING
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = footer_para.add_run(branding.WEBSITE)
    run1.font.size = Pt(branding.SIZE_HEADER_FOOTER)
    run1.font.color.rgb = NAVY
    run1.font.bold = True
    run1.font.name = branding.FONT_BODY
    run2 = footer_para.add_run(f"  |  {branding.TAGLINE}")
    run2.font.size = Pt(branding.SIZE_HEADER_FOOTER)
    run2.font.color.rgb = GRAY
    run2.font.italic = True
    run2.font.name = branding.FONT_BODY

    # Title page extraction
    title_added = False
    subtitle_added = False
    metadata_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# ') and not line.startswith('## '):
            if not title_added:
                doc.add_paragraph(line[2:].strip(), style='TGS Title')
                title_added = True
            elif not subtitle_added:
                doc.add_paragraph(line[2:].strip(), style='TGS Subtitle')
                subtitle_added = True
            i += 1
            continue
        elif line.startswith('**') and ':**' in line and not line.startswith('**URGENT'):
            metadata_lines.append(line)
            i += 1
            continue
        elif line == '---' and title_added:
            for meta_line in metadata_lines:
                para = doc.add_paragraph(style='TGS Meta')
                match = re.match(r'\*\*([^*]+)\*\*\s*(.*)', meta_line)
                if match:
                    run = para.add_run(match.group(1) + ' ')
                    run.bold = True
                    run.font.color.rgb = CHARCOAL
                    run.font.name = branding.FONT_BODY
                    run.font.size = Pt(branding.SIZE_BODY)
                    run2 = para.add_run(match.group(2))
                    run2.font.color.rgb = GRAY
                    run2.font.name = branding.FONT_BODY
                    run2.font.size = Pt(branding.SIZE_BODY)
            i += 1
            break
        elif line:
            break
        i += 1

    # Process remaining content
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Page break
        if line == '<!-- PAGE_BREAK -->':
            doc.add_page_break()
            i += 1
            continue

        # Horizontal rule
        if line == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            rows, new_i = parse_markdown_table(lines, i)
            if rows:
                add_styled_table(doc, rows)
            i = new_i
            continue

        # Blockquote (callout box)
        if line.startswith('> '):
            add_callout_box(doc, line[2:].strip())
            i += 1
            continue

        # H1 — skip (already processed as title)
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue

        # H2 → Heading 1
        if line.startswith('## '):
            h2_text = line[3:].strip()
            if h2_text == 'If You Only Do One Thing':
                doc.add_page_break()
            doc.add_paragraph(h2_text, style='Heading 1')
            i += 1
            continue

        # H3 → Heading 2
        if line.startswith('### '):
            doc.add_paragraph(line[4:].strip(), style='Heading 2')
            i += 1
            continue

        # H4 → Heading 3
        if line.startswith('#### '):
            doc.add_paragraph(line[5:].strip(), style='Heading 3')
            i += 1
            continue

        # Bold metadata lines
        if line.startswith('**') and ':**' in line and not line.startswith('**URGENT'):
            para = doc.add_paragraph(style='TGS Meta')
            match = re.match(r'\*\*([^*]+)\*\*\s*(.*)', line)
            if match:
                run = para.add_run(match.group(1) + ' ')
                run.bold = True
                run.font.color.rgb = CHARCOAL
                run.font.name = branding.FONT_BODY
                run.font.size = Pt(branding.SIZE_BODY)
                run2 = para.add_run(match.group(2))
                run2.font.color.rgb = GRAY
                run2.font.name = branding.FONT_BODY
                run2.font.size = Pt(branding.SIZE_BODY)
            i += 1
            continue

        # Status labels (URGENT, HIGH, MEDIUM)
        if line.startswith('**URGENT') or line.startswith('**HIGH') or line.startswith('**MEDIUM'):
            para = doc.add_paragraph(style='TGS Urgent')
            clean = line.replace('**', '').replace('\u2014', '-')
            run = para.add_run(clean)
            if 'URGENT' in line:
                run.font.color.rgb = ERROR
            elif 'HIGH' in line:
                run.font.color.rgb = GOLD_DARK
            else:
                run.font.color.rgb = GRAY
            run.font.bold = True
            run.font.size = Pt(branding.SIZE_BODY)
            i += 1
            continue

        # Bullet points (including nested)
        original_line = lines[i]
        bullet_match = re.match(r'^(\s*)-\s+(.*)$', original_line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            bullet_text = clean_markdown_link(bullet_match.group(2))
            indent_level = indent_spaces // 2

            if indent_level == 0:
                para = doc.add_paragraph(style='List Bullet')
            else:
                try:
                    para = doc.add_paragraph(style='List Bullet 2')
                except KeyError:
                    para = doc.add_paragraph(style='List Bullet')
                    para.paragraph_format.left_indent = Inches(0.25 * indent_level)

            add_inline_bold(para, bullet_text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            list_text = clean_markdown_link(re.sub(r'^\d+\.\s', '', line))
            para = doc.add_paragraph(style='List Number')
            add_inline_bold(para, list_text)
            i += 1
            continue

        # Score explanations
        if line.startswith('**Fit Score:**') or line.startswith('**Effort:**'):
            para = doc.add_paragraph(style='TGS Body')
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(line.replace('**', ''))
            run.font.size = Pt(branding.SIZE_CAPTION)
            run.font.italic = True
            run.font.color.rgb = GRAY
            i += 1
            continue

        # Regular paragraph
        if line:
            clean_line = clean_markdown_link(line)
            para = doc.add_paragraph(style='TGS Body')
            add_inline_bold(para, clean_line)

        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    parser = argparse.ArgumentParser(
        description='TheGrantScout Markdown to Word Converter'
    )
    parser.add_argument('--input', '-i', required=True, help='Input markdown file')
    parser.add_argument('--output', '-o', help='Output .docx file (default: same name, .docx extension)')
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    if args.output:
        docx_path = args.output
    else:
        docx_path = os.path.splitext(args.input)[0] + '.docx'

    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.isdir(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    convert_markdown_to_docx(args.input, docx_path)


if __name__ == '__main__':
    main()
